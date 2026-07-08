#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""文档编辑对外 API：导入、列表、批注、编辑、导出。"""
import os
import re
import json
import shutil
import zipfile
from datetime import datetime

from doc_paths import (
    DocDir,
    ReadJson,
    WriteJson,
    ReadManifest,
    WriteManifest,
    NewDocId,
    CalcTodoProgress,
)
from docx_parser import *  # noqa: F403
import doc_editor_html as dhtml
from doc_revisions import (
    HasWorkingStash,
    GetWorkingStatus,
    _QuickWorkingStatus,
    _LoadRevisionList,
)


def BootstrapDocPreview(sdocid):
    """批注/预览解析失败时的兜底，保证文档仍可编辑。"""
    sdir = DocDir(sdocid)
    WriteJson(os.path.join(sdir, "comments.json"), {
        "extracted": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "items": [],
        "extract_error": True,
    })
    WriteJson(os.path.join(sdir, "todos.json"), {
        "updated": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "items": [],
    })
    try:
        dhtml.BuildPreview(sdocid, [])
    except Exception:
        with open(os.path.join(sdir, "preview.html"), "w", encoding="utf-8") as f:
            f.write('<div class="docpreview"><p class="meta">预览暂不可用，可在编辑器中继续修改正文</p></div>')


def ImportDocx(bcontent, sfilename, stitle=None, vtags=None):
    sfilename = os.path.basename(sfilename)
    if not sfilename.lower().endswith(".docx"):
        raise ValueError("仅支持 .docx 文件")
    _ValidateDocxBytes(bcontent)
    bcontent = _RepairDocxBytes(bcontent)
    sdocid = NewDocId(sfilename)
    sdir = DocDir(sdocid)
    os.makedirs(sdir, exist_ok=True)
    scurrent = os.path.join(sdir, "current.docx")
    soriginal = os.path.join(sdir, "original.docx")
    with open(scurrent, "wb") as f:
        f.write(bcontent)
    shutil.copy2(scurrent, soriginal)
    from docx import Document
    try:
        Document(scurrent)
    except Exception as e:
        shutil.rmtree(sdir, ignore_errors=True)
        raise ValueError("Word 文档无法打开：%s" % e)
    stamp = datetime.now().strftime("%Y-%m-%d %H:%M")
    ometa = {
        "id": sdocid,
        "title": stitle or os.path.splitext(sfilename)[0],
        "filename": sfilename,
        "tags": vtags or [],
        "created": stamp,
        "updated": stamp,
        "revisions": [],
    }
    WriteJson(os.path.join(sdir, "meta.json"), ometa)
    omanifest = ReadManifest()
    omanifest["docs"].insert(0, {"id": sdocid, "title": ometa["title"], "updated": stamp})
    WriteManifest(omanifest)
    try:
        ExtractComments(sdocid)
    except Exception:
        BootstrapDocPreview(sdocid)
    dhtml.InvalidateEditorCache(sdocid)
    return {"id": sdocid, "title": ometa["title"]}


def ListDocs(stag=None):
    omanifest = ReadManifest()
    vresult = []
    for oitem in omanifest.get("docs", []):
        sdocid = oitem["id"]
        spath = os.path.join(DocDir(sdocid), "meta.json")
        if not os.path.isfile(spath):
            continue
        ometa = ReadJson(spath)
        if stag and stag not in ometa.get("tags", []):
            continue
        otodos = ReadJson(os.path.join(DocDir(sdocid), "todos.json"), {"items": []})
        ndone = sum(1 for x in otodos.get("items", []) if x.get("status") == "done")
        ntotal = len(otodos.get("items", []))
        odirty = _QuickWorkingStatus(sdocid)
        vresult.append({
            "id": sdocid,
            "title": ometa.get("title", sdocid),
            "tags": ometa.get("tags", []),
            "updated": ometa.get("updated", ""),
            "todo_done": ndone,
            "todo_total": ntotal,
            "progress": int(ndone * 100 / ntotal) if ntotal else 100,
            "is_dirty": odirty.get("is_dirty", False),
            "commit_count": len(_LoadRevisionList(sdocid)),
        })
    return {"docs": vresult}


def GetDocDetail(sdocid, blight=False):
    sdir = DocDir(sdocid)
    if not os.path.isdir(sdir):
        raise ValueError("文档不存在")
    ometa = ReadJson(os.path.join(sdir, "meta.json"))
    ocomments = ReadJson(os.path.join(sdir, "comments.json"), {"items": []})
    otodos = ReadJson(os.path.join(sdir, "todos.json"), {"items": []})
    vrevs = _LoadRevisionList(sdocid)
    return {
        "meta": ometa,
        "comments": ocomments.get("items", []),
        "todos": otodos.get("items", []),
        "progress": CalcTodoProgress(otodos),
        "revisions": vrevs[:20],
        "preview_ready": os.path.isfile(os.path.join(sdir, "preview.html")),
        "has_working_stash": HasWorkingStash(sdocid),
        "working_status": GetWorkingStatus(sdocid, bfull=not blight),
    }


def UpdateDocMeta(sdocid, stitle=None, vtags=None):
    spath = os.path.join(DocDir(sdocid), "meta.json")
    if not os.path.isfile(spath):
        raise ValueError("文档不存在")
    ometa = ReadJson(spath)
    if stitle:
        ometa["title"] = stitle.strip()
    if vtags is not None:
        ometa["tags"] = [t.strip() for t in vtags if t.strip()]
    ometa["updated"] = datetime.now().strftime("%Y-%m-%d %H:%M")
    WriteJson(spath, ometa)
    omanifest = ReadManifest()
    for oitem in omanifest.get("docs", []):
        if oitem["id"] == sdocid:
            oitem["title"] = ometa["title"]
            oitem["updated"] = ometa["updated"]
    WriteManifest(omanifest)
    return ometa


def GetMediaBytes(sdocid, sfname):
    sfname = os.path.basename((sfname or "").replace("\\", "/"))
    if not sfname or ".." in sfname:
        raise ValueError("无效的图片名")
    scurrent = os.path.join(DocDir(sdocid), "current.docx")
    if not os.path.isfile(scurrent):
        raise ValueError("文档不存在")
    with zipfile.ZipFile(scurrent, "r") as oz:
        for spath in ("word/media/%s" % sfname, "word/%s" % sfname):
            if spath in oz.namelist():
                bdata = oz.read(spath)
                return bdata, _MimeFromExt(os.path.splitext(sfname)[1])
    raise ValueError("图片不存在")


def ExtractComments(sdocid):
    sdir = DocDir(sdocid)
    scurrent = os.path.join(sdir, "current.docx")
    if not os.path.isfile(scurrent):
        raise ValueError("文档文件不存在")
    with zipfile.ZipFile(scurrent, "r") as oz:
        scomments = oz.read("word/comments.xml").decode("utf-8") if "word/comments.xml" in oz.namelist() else ""
        try:
            sdocument = oz.read("word/document.xml").decode("utf-8")
        except Exception:
            sdocument = ""
    try:
        vraw = _ParseCommentsXml(scomments)
    except Exception:
        vraw = []
    try:
        oparas = _MapCommentParas(sdocument)
    except Exception:
        oparas = {}
    vitems = []
    for i, oc in enumerate(vraw):
        scid = oc["id"]
        npara = oparas.get(scid, -1)
        vitems.append({
            "id": scid or str(i),
            "author": oc["author"],
            "date": oc["date"],
            "text": oc["text"],
            "para_index": npara,
            "excerpt": _ParaExcerpt(scurrent, npara) if npara >= 0 else "",
            "status": "pending",
        })
    otodos_old = ReadJson(os.path.join(sdir, "todos.json"), {"items": []})
    odone_map = {x.get("comment_id"): x.get("status") for x in otodos_old.get("items", []) if x.get("status") == "done"}
    for oc in vitems:
        if oc["id"] in odone_map:
            oc["status"] = "done"
    WriteJson(os.path.join(sdir, "comments.json"), {
        "extracted": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "items": vitems,
    })
    vtodos = []
    for oc in vitems:
        vtodos.append({
            "id": "todo-%s" % oc["id"],
            "comment_id": oc["id"],
            "text": "[%s] %s" % (oc["author"] or "批注", oc["text"]),
            "para_index": oc["para_index"],
            "status": oc["status"],
        })
    WriteJson(os.path.join(sdir, "todos.json"), {
        "updated": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "items": vtodos,
    })
    dhtml.BuildPreview(sdocid, vitems)
    ometa = ReadJson(os.path.join(sdir, "meta.json"))
    ometa["updated"] = datetime.now().strftime("%Y-%m-%d %H:%M")
    ometa["comment_count"] = len(vitems)
    WriteJson(os.path.join(sdir, "meta.json"), ometa)
    return {"comments": len(vitems), "todos": len(vtodos)}


def GetPreviewHtml(sdocid):
    spath = os.path.join(DocDir(sdocid), "preview.html")
    if not os.path.isfile(spath):
        ExtractComments(sdocid)
        spath = os.path.join(DocDir(sdocid), "preview.html")
    with open(spath, "r", encoding="utf-8") as f:
        return f.read()


def MarkTodoDone(sdocid, stodo_id, bdone=True):
    sdir = DocDir(sdocid)
    otodos = ReadJson(os.path.join(sdir, "todos.json"), {"items": []})
    ocomments = ReadJson(os.path.join(sdir, "comments.json"), {"items": []})
    sstatus = "done" if bdone else "pending"
    scid = None
    for oitem in otodos.get("items", []):
        if oitem["id"] == stodo_id:
            oitem["status"] = sstatus
            scid = oitem.get("comment_id")
            break
    for oc in ocomments.get("items", []):
        if scid and oc["id"] == scid:
            oc["status"] = sstatus
    WriteJson(os.path.join(sdir, "todos.json"), otodos)
    WriteJson(os.path.join(sdir, "comments.json"), ocomments)
    return {"progress": CalcTodoProgress(otodos)}


def SetParaPlainText(opara, stext):
    if _ParaHasDrawing(opara):
        _SetParaTextPreserveMedia(opara, stext)
        return
    _ClearParaRuns(opara)
    if stext:
        _EmitStyledText(opara, {"text": stext})


def ApplyEdit(sdocid, npara_index, snew_text, scomment_id=None, shtml=None, spara_style=None):
    try:
        from docx import Document
    except ImportError as e:
        raise ValueError("服务器缺少 python-docx 依赖，无法保存文档") from e
    scurrent = os.path.join(DocDir(sdocid), "current.docx")
    if not os.path.isfile(scurrent):
        raise ValueError("文档不存在")
    odoc = Document(scurrent)
    vparas = _BodyParagraphs(odoc)
    if npara_index < 0 or npara_index >= len(vparas):
        raise ValueError("段落索引无效")
    opara = vparas[npara_index]
    if spara_style:
        _ApplyParaLayout(opara, spara_style)
    if shtml and shtml.strip():
        _ApplyRichHtmlToPara(opara, shtml)
        stext = _SanitizeParaText(opara.text or "")
    else:
        stext = _SanitizeParaText(snew_text)
        SetParaPlainText(opara, stext)
    odoc.save(scurrent)
    dhtml.InvalidateEditorCache(sdocid)
    dhtml.BuildPreview(sdocid)
    ometa = ReadJson(os.path.join(DocDir(sdocid), "meta.json"))
    ometa["updated"] = datetime.now().strftime("%Y-%m-%d %H:%M")
    WriteJson(os.path.join(DocDir(sdocid), "meta.json"), ometa)
    if scomment_id:
        MarkTodoDone(sdocid, "todo-%s" % scomment_id, True)
    return {"para_index": npara_index}


def NormalizeExportFilename(sfilename):
    sfilename = os.path.basename((sfilename or "export.docx").strip())
    if not sfilename.lower().endswith(".docx"):
        sfilename += ".docx"
    sfilename = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", sfilename).strip("._") or "export.docx"
    return sfilename


def BuildExportDocx(sdocid):
    scurrent = os.path.join(DocDir(sdocid), "current.docx")
    if not os.path.isfile(scurrent):
        raise ValueError("文档不存在")
    try:
        from docx import Document
    except ImportError as e:
        raise ValueError("服务器缺少 python-docx 依赖，无法导出 Word 文档") from e
    odoc = Document(scurrent)
    _CleanDocxUiArtifacts(odoc)
    return odoc


def ExportDocBytes(sdocid, sfilename):
    import io
    sfilename = NormalizeExportFilename(sfilename)
    obuf = io.BytesIO()
    BuildExportDocx(sdocid).save(obuf)
    return obuf.getvalue(), sfilename


def ExportDoc(sdocid, sdest_dir, sfilename):
    sdest_dir = os.path.expanduser((sdest_dir or "").strip())
    if not sdest_dir:
        raise ValueError("请选择有效的导出文件夹")
    sdest_dir = os.path.abspath(sdest_dir)
    os.makedirs(sdest_dir, exist_ok=True)
    sfilename = NormalizeExportFilename(sfilename)
    sdest = os.path.join(sdest_dir, sfilename)
    BuildExportDocx(sdocid).save(sdest)
    return {"path": sdest, "filename": sfilename, "dir": sdest_dir}


def DeleteDoc(sdocid):
    sdir = DocDir(sdocid)
    if os.path.isdir(sdir):
        shutil.rmtree(sdir)
    omanifest = ReadManifest()
    omanifest["docs"] = [d for d in omanifest.get("docs", []) if d["id"] != sdocid]
    WriteManifest(omanifest)
    return {"id": sdocid}
