#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""纯 .docx / OOXML 读写与 HTML 转换（doc_editor 拆分底层）。"""
import os
import re
import json
import zipfile
import urllib.parse
import xml.etree.ElementTree as ET
from html.parser import HTMLParser

WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": WNS}
RNS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
ANS = "http://schemas.openxmlformats.org/drawingml/2006/main"
RELNS = "http://schemas.openxmlformats.org/package/2006/relationships"


def _ValidateDocxBytes(bcontent):
    import io
    if not bcontent:
        raise ValueError("文件内容为空")
    try:
        with zipfile.ZipFile(io.BytesIO(bcontent), "r") as oz:
            if "word/document.xml" not in oz.namelist():
                raise ValueError("不是有效的 docx 文件（缺少 document.xml）")
            oz.read("word/document.xml")
    except zipfile.BadZipFile:
        raise ValueError("不是有效的 docx 文件（若为 .doc 旧格式，请先在 Word/WPS 中另存为 .docx）")


def _XmlParseOk(bdata):
    if not bdata:
        return True
    try:
        ET.fromstring(bdata)
        return True
    except Exception:
        return False


def _StripBrokenRels(brels, vskip_targets):
    if not vskip_targets:
        return brels
    try:
        oroot = ET.fromstring(brels)
    except Exception:
        return brels
    vremove = []
    for orel in list(oroot):
        starget = (orel.get("Target") or "").replace("\\", "/")
        if not starget:
            continue
        for sskip in vskip_targets:
            if starget.endswith(sskip) or starget.endswith(os.path.basename(sskip)):
                vremove.append(orel)
                break
    for orel in vremove:
        oroot.remove(orel)
    return ET.tostring(oroot, encoding="utf-8", xml_declaration=True)


def _RepairDocxBytes(bcontent):
    """剥离损坏的可选 XML（如 WPS 批注），避免 python-docx 无法打开。"""
    import io
    voptional_keys = ("comment", "footnote", "endnote", "people", "chart")
    vskip = set()
    vin = io.BytesIO(bcontent)
    with zipfile.ZipFile(vin, "r") as zin:
        for sname in zin.namelist():
            if not sname.startswith("word/") or not sname.endswith(".xml"):
                continue
            if sname == "word/document.xml":
                if not _XmlParseOk(zin.read(sname)):
                    raise ValueError("文档主体损坏，无法导入")
                continue
            if not any(skey in sname.lower() for skey in voptional_keys):
                continue
            if not _XmlParseOk(zin.read(sname)):
                vskip.add(sname)
        if not vskip:
            return bcontent
        vout = io.BytesIO()
        vin.seek(0)
        with zipfile.ZipFile(vin, "r") as zin2:
            with zipfile.ZipFile(vout, "w", compression=zipfile.ZIP_DEFLATED) as zout:
                for item in zin2.infolist():
                    if item.filename in vskip:
                        continue
                    data = zin2.read(item.filename)
                    if item.filename == "word/_rels/document.xml.rels":
                        data = _StripBrokenRels(data, vskip)
                    zout.writestr(item, data)
        return vout.getvalue()

def _TextFromXml(onode):
    vparts = []
    for telem in onode.iter("{%s}t" % WNS):
        if telem.text:
            vparts.append(telem.text)
    return "".join(vparts).strip()


def _ParseCommentsXml(scomments_xml):
    vcomments = []
    if not scomments_xml:
        return vcomments
    oroot = ET.fromstring(scomments_xml)
    for ocomment in oroot.findall("w:comment", NS):
        scid = ocomment.get("{%s}id" % WNS, "")
        sauthor = ocomment.get("{%s}author" % WNS, "")
        sdate = ocomment.get("{%s}date" % WNS, "")
        stext = _TextFromXml(ocomment)
        vcomments.append({
            "id": scid,
            "author": sauthor,
            "date": sdate[:19] if sdate else "",
            "text": stext,
        })
    return vcomments


def _MapCommentParas(sdocument_xml):
    """批注 id → 段落索引。"""
    if not sdocument_xml:
        return {}
    oroot = ET.fromstring(sdocument_xml)
    obody = oroot.find("w:body", NS)
    if obody is None:
        return {}
    omap = {}
    npara = -1
    for ochild in obody:
        stag = ochild.tag.split("}")[-1] if "}" in ochild.tag else ochild.tag
        if stag == "p":
            npara += 1
        for ostart in ochild.iter("{%s}commentRangeStart" % WNS):
            scid = ostart.get("{%s}id" % WNS)
            if scid is not None:
                omap[scid] = npara
    return omap


def _EscHtml(stext):
    return (stext or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")


def _EscAttr(stext):
    return _EscHtml(stext).replace("'", "&#39;")


def _MimeFromExt(sext):
    return {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".gif": "image/gif",
        ".bmp": "image/bmp",
        ".webp": "image/webp",
        ".emf": "image/emf",
        ".wmf": "image/wmf",
    }.get(sext.lower(), "application/octet-stream")


def _LoadImageRels(scurrent):
    with zipfile.ZipFile(scurrent, "r") as oz:
        srels = "word/_rels/document.xml.rels"
        if srels not in oz.namelist():
            return {}
        oroot = ET.fromstring(oz.read(srels))
        omap = {}
        for orel in oroot:
            if orel.get("Type", "").endswith("/image"):
                omap[orel.get("Id")] = orel.get("Target", "")
        return omap


def _LoadImageUrls(sdocid, scurrent):
    """rId / 文件名 → 媒体接口 URL，避免整页内嵌 base64。"""
    orels = _LoadImageRels(scurrent)
    odata = {}
    for srid, starget in orels.items():
        sfname = os.path.basename(starget.replace("\\", "/"))
        if not sfname:
            continue
        surl = "/api/docs/media?id=%s&file=%s" % (
            urllib.parse.quote(sdocid, safe=""),
            urllib.parse.quote(sfname, safe=""),
        )
        odata[srid] = surl
        odata[sfname] = surl
    return odata


def _ParaHasDrawing(opara):
    for orun in opara.runs:
        if orun._element.findall(".//{%s}drawing" % WNS):
            return True
    return False


from docx_style import (
    _RgbToHex,
    _TwipsToPt,
    _ElemRPr,
    _ElemPPr,
    _LoadThemeFonts,
    _ResolveThemeFontName,
    _RFontsFromRPr,
    _SzPtFromRPr,
    _ColorFromRPr,
    _FlagFromRPr,
    _JcValToCss,
    _AlignEnumToCss,
    _IndLayoutFromPPr,
    _SpacingLayoutFromPPr,
    _StylePPrRPr,
    _ParaLayoutDict,
    _LayoutDictToCss,
    _ParaStyleCss,
)

def _RunFormatDict(orun, opara=None, othemes=None):
    ofmt = {
        "bold": bool(orun.bold),
        "italic": bool(orun.italic),
        "underline": bool(orun.underline),
        "font": (orun.font.name or "").strip(),
        "size_pt": None,
        "color": "",
    }
    if orun.font.size:
        try:
            ofmt["size_pt"] = round(orun.font.size.pt, 1)
        except Exception:
            pass
    if orun.font.color and orun.font.color.rgb:
        ofmt["color"] = _RgbToHex(orun.font.color.rgb)

    rr = _ElemRPr(orun._element)
    if rr is not None:
        if not ofmt["bold"] and _FlagFromRPr(rr, "b"):
            ofmt["bold"] = True
        if not ofmt["italic"] and _FlagFromRPr(rr, "i"):
            ofmt["italic"] = True
        if not ofmt["underline"] and _FlagFromRPr(rr, "u"):
            ofmt["underline"] = True
        if not ofmt["font"]:
            ofmt["font"] = _RFontsFromRPr(rr, othemes)
        if ofmt["size_pt"] is None:
            ofmt["size_pt"] = _SzPtFromRPr(rr)
        if not ofmt["color"]:
            ofmt["color"] = _ColorFromRPr(rr)

    ppr = _ElemPPr(opara._element) if opara is not None else None
    pr = _ElemRPr(ppr) if ppr is not None else None
    if pr is not None:
        if not ofmt["font"]:
            ofmt["font"] = _RFontsFromRPr(pr, othemes)
        if ofmt["size_pt"] is None:
            ofmt["size_pt"] = _SzPtFromRPr(pr)
        if not ofmt["color"]:
            ofmt["color"] = _ColorFromRPr(pr)
        if not ofmt["bold"] and _FlagFromRPr(pr, "b"):
            ofmt["bold"] = True
        if not ofmt["italic"] and _FlagFromRPr(pr, "i"):
            ofmt["italic"] = True
        if not ofmt["underline"] and _FlagFromRPr(pr, "u"):
            ofmt["underline"] = True

    _, sr = _StylePPrRPr(opara) if opara is not None else (None, None)
    if sr is not None:
        if not ofmt["font"]:
            ofmt["font"] = _RFontsFromRPr(sr, othemes)
        if ofmt["size_pt"] is None:
            ofmt["size_pt"] = _SzPtFromRPr(sr)
        if not ofmt["color"]:
            ofmt["color"] = _ColorFromRPr(sr)
        if not ofmt["bold"] and _FlagFromRPr(sr, "b"):
            ofmt["bold"] = True
        if not ofmt["italic"] and _FlagFromRPr(sr, "i"):
            ofmt["italic"] = True
        if not ofmt["underline"] and _FlagFromRPr(sr, "u"):
            ofmt["underline"] = True

    if opara is not None:
        try:
            if opara.style and opara.style.font:
                if not ofmt["font"] and opara.style.font.name:
                    ofmt["font"] = opara.style.font.name.strip()
                if ofmt["size_pt"] is None and opara.style.font.size:
                    ofmt["size_pt"] = round(opara.style.font.size.pt, 1)
        except Exception:
            pass
    return ofmt


def _FormatDictToCss(ofmt):
    vstyles = []
    if ofmt.get("font"):
        vstyles.append("font-family:%s" % ofmt["font"])
    if ofmt.get("size_pt"):
        vstyles.append("font-size:%gpt" % ofmt["size_pt"])
    if ofmt.get("color"):
        vstyles.append("color:%s" % ofmt["color"])
    return ";".join(vstyles)


def _ApplyInlineFmtToHtml(shtml, ofmt):
    if ofmt.get("bold"):
        shtml = "<strong>%s</strong>" % shtml
    if ofmt.get("italic"):
        shtml = "<em>%s</em>" % shtml
    if ofmt.get("underline"):
        shtml = "<u>%s</u>" % shtml
    sstyle = _FormatDictToCss(ofmt)
    if sstyle:
        shtml = '<span style="%s">%s</span>' % (_EscAttr(sstyle), shtml)
    return shtml


def _WrapText(stext, orun, opara=None, othemes=None):
    if not stext:
        return ""
    ofmt = _RunFormatDict(orun, opara, othemes)
    vlines = stext.split("\n")
    vparts = []
    for i, sline in enumerate(vlines):
        if sline:
            vparts.append(_ApplyInlineFmtToHtml(_EscHtml(sline), ofmt))
        if i < len(vlines) - 1:
            vparts.append("<br>")
    return "".join(vparts)


def _RunToHtml(orun, oimgurls, beditable=False, opara=None, othemes=None):
    vparts = []
    bhas_child = False
    for ochild in orun._element:
        stag = ochild.tag.split("}")[-1] if "}" in ochild.tag else ochild.tag
        if stag == "t":
            bhas_child = True
            st = ochild.text or ""
            if st:
                vparts.append(_WrapText(st, orun, opara, othemes))
        elif stag == "br":
            bhas_child = True
            vparts.append("<br>")
        elif stag == "tab":
            bhas_child = True
            vparts.append("&#9;")
        elif stag == "drawing":
            bhas_child = True
            for oblip in ochild.iter("{%s}blip" % ANS):
                srid = oblip.get("{%s}embed" % RNS)
                surl = oimgurls.get(srid, "")
                if surl:
                    sattr = ' contenteditable="false"' if beditable else ""
                    vparts.append(
                        '<span class="imgwrap"%s><img class="docimg" src="%s" alt=""></span>'
                        % (sattr, surl)
                    )
    if not bhas_child and orun.text:
        vparts.append(_WrapText(orun.text, orun, opara, othemes))
    if not vparts:
        for oblip in orun._element.iter("{%s}blip" % ANS):
            srid = oblip.get("{%s}embed" % RNS)
            surl = oimgurls.get(srid, "")
            if surl:
                sattr = ' contenteditable="false"' if beditable else ""
                vparts.append(
                    '<span class="imgwrap"%s><img class="docimg" src="%s" alt=""></span>'
                    % (sattr, surl)
                )
    return "".join(vparts)


def _PlainTextToHtml(stext):
    if not stext:
        return "&#160;"
    vlines = stext.split("\n")
    vparts = []
    for i, sline in enumerate(vlines):
        if sline:
            vparts.append(_EscHtml(sline))
        if i < len(vlines) - 1:
            vparts.append("<br>")
    return "".join(vparts) or "&#160;"


def _ParaToHtml(opara, oimgurls, beditable=False, othemes=None):
    if not opara.runs:
        sbody = _PlainTextToHtml(opara.text or "")
        sstyle = _ParaStyleCss(opara, othemes)
        if sstyle and sbody != "&#160;":
            return '<span style="%s">%s</span>' % (_EscAttr(sstyle), sbody)
        return sbody
    sbody = "".join(_RunToHtml(r, oimgurls, beditable, opara, othemes) for r in opara.runs)
    return sbody or "&#160;"


def _TableToHtml(otable, oimgurls, beditable=False, othemes=None):
    vrows = []
    for orow in otable.rows:
        vcells = []
        for ocell in orow.cells:
            vparts = []
            for opara in ocell.paragraphs:
                vparts.append(_ParaToHtml(opara, oimgurls, beditable, othemes))
            vcells.append("<td>%s</td>" % ("<br>".join(vparts) or "&#160;"))
        vrows.append("<tr>%s</tr>" % "".join(vcells))
    return '<table class="doctable">%s</table>' % "".join(vrows)


def _CommentMarks(npara, opara_comments):
    if npara not in opara_comments:
        return ""
    return "".join(
        '<span class="cmtmark" data-cid="%s" title="%s">📝</span>' % (
            c["id"],
            (c["text"][:40] + "…").replace('"', "'"),
        )
        for c in opara_comments[npara]
    )


_RE_LEAKED_TAG = re.compile(
    r"</?\s*(?:span|font|strong|em|u|br|p|div|b|i)\b[^<>]*>"
    r"|(?:span|font)\b[^<>]*?(?:style|face|color|size)\s*=[^<>]*>",
    re.I,
)


def _StripHtmlTags(stext):
    """剥离误写入正文的 HTML 标签（如 round-trip 泄漏的 <span style=...>）。"""
    if not stext or "<" not in stext and ">" not in stext:
        return stext or ""
    return _RE_LEAKED_TAG.sub("", stext)


def _SanitizeParaText(stext):
    """去除预览 UI 泄漏的批注图标及误写入的标签，不写入 docx。"""
    if not stext:
        return ""
    stext = stext.replace("\U0001f4dd", "").replace("📝", "")
    return _StripHtmlTags(stext).strip()


_FONT_SIZE_LEGACY = {1: 8, 2: 10, 3: 12, 4: 14, 5: 18, 6: 24, 7: 36}


def _ParseParaCss(sstyle):
    olayout = {}
    if not sstyle:
        return olayout
    for spart in sstyle.split(";"):
        if ":" not in spart:
            continue
        skey, sval = spart.split(":", 1)
        skey = skey.strip().lower()
        sval = sval.strip().strip("'\"")
        if skey == "text-align":
            olayout["align"] = sval
        elif skey == "margin-left":
            npt = _ParseFontSize(sval)
            if npt is not None:
                olayout["left_pt"] = npt
        elif skey == "margin-right":
            npt = _ParseFontSize(sval)
            if npt is not None:
                olayout["right_pt"] = npt
        elif skey == "text-indent":
            npt = _ParseFontSize(sval)
            if npt is not None:
                if npt < 0:
                    olayout["hang_pt"] = abs(npt)
                else:
                    olayout["first_pt"] = npt
        elif skey == "padding-left":
            npt = _ParseFontSize(sval)
            if npt is not None:
                olayout["left_pt"] = npt
        elif skey == "margin-top":
            npt = _ParseFontSize(sval)
            if npt is not None:
                olayout["space_before_pt"] = npt
        elif skey == "margin-bottom":
            npt = _ParseFontSize(sval)
            if npt is not None:
                olayout["space_after_pt"] = npt
        elif skey == "line-height":
            try:
                if sval.endswith("pt"):
                    olayout["line_height_pt"] = _ParseFontSize(sval)
                else:
                    olayout["line_height"] = round(float(sval), 2)
            except Exception:
                pass
    return olayout


def _ApplyParaLayout(opara, scss):
    olayout = _ParseParaCss(scss)
    if not olayout:
        return
    pf = opara.paragraph_format
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    from docx.enum.text import WD_LINE_SPACING
    omap = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    if "align" in olayout:
        pf.alignment = omap.get(olayout["align"], WD_ALIGN_PARAGRAPH.LEFT)
    if "left_pt" in olayout:
        pf.left_indent = Pt(olayout["left_pt"])
    if "right_pt" in olayout:
        pf.right_indent = Pt(olayout["right_pt"])
    if "hang_pt" in olayout:
        pf.first_line_indent = Pt(-olayout["hang_pt"])
    elif "first_pt" in olayout:
        pf.first_line_indent = Pt(olayout["first_pt"])
    if "space_before_pt" in olayout:
        pf.space_before = Pt(olayout["space_before_pt"])
    if "space_after_pt" in olayout:
        pf.space_after = Pt(olayout["space_after_pt"])
    if "line_height" in olayout:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = olayout["line_height"]
    elif "line_height_pt" in olayout:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(olayout["line_height_pt"])


def _ParseCssStyle(sstyle):
    ostyle = {}
    if not sstyle:
        return ostyle
    for spart in sstyle.split(";"):
        if ":" not in spart:
            continue
        skey, sval = spart.split(":", 1)
        skey = skey.strip().lower()
        sval = sval.strip().strip("'\"")
        if skey == "font-family":
            sval = sval.split(",")[0].strip()
            if sval:
                ostyle["font"] = sval
        elif skey == "font-size":
            osize = _ParseFontSize(sval)
            if osize:
                ostyle["size_pt"] = osize
        elif skey == "color":
            orgb = _ParseColor(sval)
            if orgb:
                ostyle["color"] = orgb
        elif skey == "font-weight" and sval in ("bold", "700", "800", "900"):
            ostyle["bold"] = True
        elif skey == "font-style" and sval == "italic":
            ostyle["italic"] = True
        elif skey == "text-decoration" and "underline" in sval:
            ostyle["underline"] = True
    return ostyle


def _ParseFontSize(sval):
    if not sval:
        return None
    sm = re.match(r"^([\d.]+)\s*(pt|px)?$", sval.lower())
    if not sm:
        return None
    nval = float(sm.group(1))
    sunit = sm.group(2) or "pt"
    if sunit == "px":
        nval = nval * 0.75
    return round(nval, 1) if nval > 0 else None


def _ParseColor(sval):
    if not sval:
        return None
    sval = sval.strip().lower()
    sm = re.match(r"^#([0-9a-f]{6})$", sval)
    if sm:
        shex = sm.group(1)
        return (int(shex[0:2], 16), int(shex[2:4], 16), int(shex[4:6], 16))
    sm = re.match(r"^rgb\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)$", sval)
    if sm:
        return (int(sm.group(1)), int(sm.group(2)), int(sm.group(3)))
    omap = {
        "black": (0, 0, 0), "white": (255, 255, 255), "red": (255, 0, 0),
        "blue": (0, 0, 255), "green": (0, 128, 0), "gray": (128, 128, 128),
        "grey": (128, 128, 128),
    }
    return omap.get(sval)


def _MergeStyle(obase, oextra):
    oout = dict(obase or {})
    for skey, sval in (oextra or {}).items():
        if sval is not None:
            oout[skey] = sval
    return oout


class _InlineHtmlParser(HTMLParser):
    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.vsegments = []
        self.vstack = [{}]

    def _PushStyle(self, oextra):
        self.vstack.append(_MergeStyle(self.vstack[-1], oextra))

    def handle_starttag(self, stag, vattrs):
        stag = stag.lower()
        oattrs = dict(vattrs)
        if stag in ("b", "strong"):
            self._PushStyle({"bold": True})
        elif stag in ("i", "em"):
            self._PushStyle({"italic": True})
        elif stag == "u":
            self._PushStyle({"underline": True})
        elif stag == "span":
            self._PushStyle(_ParseCssStyle(oattrs.get("style", "")))
        elif stag == "font":
            oextra = {}
            if oattrs.get("face"):
                oextra["font"] = oattrs.get("face").split(",")[0].strip()
            if oattrs.get("color"):
                orgb = _ParseColor(oattrs.get("color"))
                if orgb:
                    oextra["color"] = orgb
            if oattrs.get("size"):
                try:
                    nlegacy = _FONT_SIZE_LEGACY.get(int(oattrs.get("size")))
                    if nlegacy:
                        oextra["size_pt"] = float(nlegacy)
                except Exception:
                    pass
            self._PushStyle(oextra)
        elif stag == "br":
            self.vsegments.append(_MergeStyle(self.vstack[-1], {"text": "\n"}))

    def handle_endtag(self, stag):
        stag = stag.lower()
        if stag in ("b", "strong", "i", "em", "u", "span", "font") and len(self.vstack) > 1:
            self.vstack.pop()

    def handle_data(self, sdata):
        if not sdata:
            return
        self.vsegments.append(_MergeStyle(self.vstack[-1], {"text": sdata}))


def _NormalizeEditorHtml(shtml):
    if not shtml:
        return ""
    shtml = shtml.strip()
    shtml = re.sub(r"</div>\s*<div\b[^>]*>", "<br>", shtml, flags=re.I)
    shtml = re.sub(r"</p>\s*<p\b[^>]*>", "<br>", shtml, flags=re.I)
    shtml = re.sub(r"<div\b[^>]*>", "", shtml, flags=re.I)
    shtml = re.sub(r"</div>", "<br>", shtml, flags=re.I)
    shtml = re.sub(r"<p\b[^>]*>", "", shtml, flags=re.I)
    shtml = re.sub(r"</p>", "<br>", shtml, flags=re.I)
    shtml = re.sub(r"<br\s*/?>", "<br>", shtml, flags=re.I)
    shtml = re.sub(r"(<br>){2,}", "<br>", shtml, flags=re.I)
    return shtml.strip().strip("<br>")


def _ParseInlineHtml(shtml):
    sparser = _InlineHtmlParser()
    try:
        sparser.feed(_NormalizeEditorHtml(shtml))
        sparser.close()
    except Exception:
        return [{"text": _SanitizeParaText(re.sub(r"<[^>]+>", "", shtml or ""))}]
    vmerged = []
    for oseg in sparser.vsegments:
        stext = _StripHtmlTags(oseg.get("text", ""))
        if not stext:
            continue
        ostyle = {k: v for k, v in oseg.items() if k != "text"}
        if vmerged:
            oprev = {k: v for k, v in vmerged[-1].items() if k != "text"}
            if oprev == ostyle:
                vmerged[-1]["text"] += stext
                continue
        vmerged.append(dict(ostyle, text=stext))
    return vmerged or [{"text": ""}]


def _ClearParaRuns(opara):
    for orun in list(opara.runs):
        orun._element.getparent().remove(orun._element)


def _ApplySegmentStyle(orun, oseg):
    if oseg.get("bold"):
        orun.bold = True
    if oseg.get("italic"):
        orun.italic = True
    if oseg.get("underline"):
        orun.underline = True
    if oseg.get("font"):
        orun.font.name = oseg["font"]
    if oseg.get("size_pt"):
        from docx.shared import Pt
        orun.font.size = Pt(oseg["size_pt"])
    if oseg.get("color"):
        from docx.shared import RGBColor
        orgb = oseg["color"]
        orun.font.color.rgb = RGBColor(orgb[0], orgb[1], orgb[2])


def _EmitStyledText(opara, oseg):
    stext = oseg.get("text", "")
    if not stext:
        return
    vparts = stext.split("\n")
    for i, sline in enumerate(vparts):
        if sline:
            orun = opara.add_run(sline)
            _ApplySegmentStyle(orun, oseg)
        if i < len(vparts) - 1:
            orun = opara.add_run("")
            _ApplySegmentStyle(orun, oseg)
            orun.add_break()


def _ApplyRichHtmlToPara(opara, shtml):
    vsegments = _ParseInlineHtml(shtml)
    if _ParaHasDrawing(opara):
        stext = _SanitizeParaText("".join(s.get("text", "") for s in vsegments))
        _SetParaTextPreserveMedia(opara, stext)
        return
    _ClearParaRuns(opara)
    for oseg in vsegments:
        _EmitStyledText(opara, oseg)


def _ParaComparable(opara, othemes=None):
    vruns = []
    for orun in opara.runs:
        ofmt = _RunFormatDict(orun, opara, othemes)
        vruns.append({
            "t": _StripHtmlTags(orun.text or ""),
            "b": ofmt.get("bold"),
            "i": ofmt.get("italic"),
            "u": ofmt.get("underline"),
            "f": ofmt.get("font") or "",
            "s": ofmt.get("size_pt") or "",
            "c": ofmt.get("color") or "",
        })
    return json.dumps({
        "layout": _ParaLayoutDict(opara, othemes),
        "runs": vruns,
    }, ensure_ascii=False)


def _BodyParagraphs(odoc):
    """与编辑器 para index 一致：仅 body 直下段落（跳过表格内段落）。"""
    from docx.text.paragraph import Paragraph
    vparas = []
    for ochild in odoc.element.body:
        stag = ochild.tag.split("}")[-1] if "}" in ochild.tag else ochild.tag
        if stag == "p":
            vparas.append(Paragraph(ochild, odoc))
    return vparas


def _IterAllParagraphs(odoc):
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    for ochild in odoc.element.body:
        stag = ochild.tag.split("}")[-1] if "}" in ochild.tag else ochild.tag
        if stag == "p":
            yield Paragraph(ochild, odoc)
        elif stag == "tbl":
            otable = Table(ochild, odoc)
            for orow in otable.rows:
                for ocell in orow.cells:
                    for opara in ocell.paragraphs:
                        yield opara


def _CleanDocxUiArtifacts(odoc):
    bchanged = False
    for opara in _IterAllParagraphs(odoc):
        stext = opara.text or ""
        sclean = _SanitizeParaText(stext)
        if sclean == stext:
            continue
        bchanged = True
        if _ParaHasDrawing(opara):
            _SetParaTextPreserveMedia(opara, sclean)
        else:
            opara.text = sclean
    return bchanged


def _ParaBlockHtml(npara, opara, oimgurls, opara_comments_all, opara_comments_pending, beditable=False, othemes=None):
    sbody = _ParaToHtml(opara, oimgurls, beditable, othemes)
    scls = "docpara"
    if beditable:
        scls += " docpara-editable"
    if npara in opara_comments_pending:
        scls += " has-comment"
    if _ParaHasDrawing(opara):
        scls += " has-image"
    sedit = ' contenteditable="true"' if beditable else ""
    splain = _EscAttr(_SanitizeParaText(opara.text or ""))
    spstyle = _ParaStyleCss(opara, othemes)
    sstyleattr = (' style="%s"' % _EscAttr(spstyle)) if spstyle else ""
    spstyleattr = (' data-pstyle="%s"' % _EscAttr(spstyle)) if spstyle else ""
    smarks = _CommentMarks(npara, opara_comments_all)
    if beditable and smarks:
        return (
            '<div class="parablock" data-para="%d">'
            '<div class="cmtmarks" contenteditable="false">%s</div>'
            '<p id="para-%d" class="%s" data-para="%d" data-plain="%s"%s%s%s>%s</p>'
            '</div>'
        ) % (npara, smarks, npara, scls, npara, splain, spstyleattr, sstyleattr, sedit, sbody)
    return (
        '<p id="para-%d" class="%s" data-para="%d" data-plain="%s"%s%s%s>%s%s</p>'
        % (npara, scls, npara, splain, spstyleattr, sstyleattr, sedit, smarks, sbody)
    )


def _BuildDocumentBodyHtml(odoc, oimgurls, opara_comments_all, opara_comments_pending, beditable=False, othemes=None):
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    vhtml = []
    npara = -1
    for ochild in odoc.element.body:
        stag = ochild.tag.split("}")[-1] if "}" in ochild.tag else ochild.tag
        if stag == "p":
            npara += 1
            opara = Paragraph(ochild, odoc)
            vhtml.append(_ParaBlockHtml(
                npara, opara, oimgurls, opara_comments_all, opara_comments_pending, beditable, othemes,
            ))
        elif stag == "tbl":
            vhtml.append(_TableToHtml(Table(ochild, odoc), oimgurls, beditable, othemes))
    return "\n".join(vhtml)


def _SetParaTextPreserveMedia(opara, stext):
    vtext_runs = []
    for orun in opara.runs:
        if not orun._element.findall(".//{%s}drawing" % WNS):
            vtext_runs.append(orun)
    if not vtext_runs:
        if stext:
            opara.add_run(stext)
        return
    vtext_runs[0].text = stext
    for orun in vtext_runs[1:]:
        oparent = orun._element.getparent()
        if oparent is not None:
            oparent.remove(orun._element)

def _ParaExcerpt(sdocx_path, nindex):
    from docx import Document
    odoc = Document(sdocx_path)
    if nindex < 0 or nindex >= len(odoc.paragraphs):
        return ""
    return (odoc.paragraphs[nindex].text or "")[:120]

