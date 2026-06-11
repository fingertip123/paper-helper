#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""文档编辑：docx 导入、批注抓取、Todo、修订留痕（WPS + Git 范式）。"""
import os
import re
import json
import shutil
import base64
import hashlib
import zipfile
import urllib.parse
import xml.etree.ElementTree as ET
from datetime import datetime
from html.parser import HTMLParser

import topic_manager as topics

WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": WNS}
RNS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
ANS = "http://schemas.openxmlformats.org/drawingml/2006/main"
RELNS = "http://schemas.openxmlformats.org/package/2006/relationships"

_docsdir = ""
_oEditorHtmlCache = {}
_nEditorCacheMax = 10


def Init(ntopicdir):
    global _docsdir
    _docsdir = os.path.join(ntopicdir, "docs")
    os.makedirs(_docsdir, exist_ok=True)


def DocsDir():
    return _docsdir


def ManifestPath():
    return os.path.join(_docsdir, "index.json")


def DocDir(sdocid):
    return os.path.join(_docsdir, sdocid)


def ReadJson(spath, sdefault=None):
    if not os.path.isfile(spath):
        return sdefault if sdefault is not None else {}
    with open(spath, "r", encoding="utf-8") as f:
        return json.load(f)


def WriteJson(spath, odata):
    os.makedirs(os.path.dirname(spath), exist_ok=True)
    with open(spath, "w", encoding="utf-8") as f:
        json.dump(odata, f, ensure_ascii=False, indent=2)


def ReadManifest():
    omanifest = ReadJson(ManifestPath(), {"docs": []})
    if "docs" not in omanifest:
        omanifest["docs"] = []
    return omanifest


def WriteManifest(omanifest):
    WriteJson(ManifestPath(), omanifest)


def NewDocId(sfilename):
    sbase = re.sub(r"[^\w\u4e00-\u9fff-]+", "-", os.path.splitext(sfilename)[0]).strip("-").lower()
    if not sbase:
        sbase = "doc"
    sdocid = sbase[:40]
    if os.path.isdir(DocDir(sdocid)):
        sdocid = "%s-%s" % (sdocid[:30], datetime.now().strftime("%H%M%S"))
    return sdocid


def ImportDocx(bcontent, sfilename, stitle=None, vtags=None):
    sfilename = os.path.basename(sfilename)
    if not sfilename.lower().endswith(".docx"):
        raise ValueError("仅支持 .docx 文件")
    sdocid = NewDocId(sfilename)
    sdir = DocDir(sdocid)
    os.makedirs(sdir, exist_ok=True)
    scurrent = os.path.join(sdir, "current.docx")
    soriginal = os.path.join(sdir, "original.docx")
    with open(scurrent, "wb") as f:
        f.write(bcontent)
    shutil.copy2(scurrent, soriginal)
    stamp = datetime.now().strftime("%Y-%m-%d %H:%M")
    ometa = {
        "id": sdocid,
        "title": stitle or os.path.splitext(sfilename)[0],
        "filename": sfilename,
        "tags": vtags or [],
        "created": stamp,
        "updated": stamp,
        "revisions": [],
    }
    WriteJson(os.path.join(sdir, "meta.json"), ometa)
    omanifest = ReadManifest()
    omanifest["docs"].insert(0, {"id": sdocid, "title": ometa["title"], "updated": stamp})
    WriteManifest(omanifest)
    ExtractComments(sdocid)
    InvalidateEditorCache(sdocid)
    return {"id": sdocid, "title": ometa["title"]}


def ListDocs(stag=None):
    omanifest = ReadManifest()
    vresult = []
    for oitem in omanifest.get("docs", []):
        sdocid = oitem["id"]
        spath = os.path.join(DocDir(sdocid), "meta.json")
        if not os.path.isfile(spath):
            continue
        ometa = ReadJson(spath)
        if stag and stag not in ometa.get("tags", []):
            continue
        otodos = ReadJson(os.path.join(DocDir(sdocid), "todos.json"), {"items": []})
        ndone = sum(1 for x in otodos.get("items", []) if x.get("status") == "done")
        ntotal = len(otodos.get("items", []))
        odirty = _QuickWorkingStatus(sdocid)
        vresult.append({
            "id": sdocid,
            "title": ometa.get("title", sdocid),
            "tags": ometa.get("tags", []),
            "updated": ometa.get("updated", ""),
            "todo_done": ndone,
            "todo_total": ntotal,
            "progress": int(ndone * 100 / ntotal) if ntotal else 100,
            "is_dirty": odirty.get("is_dirty", False),
            "commit_count": len(_LoadRevisionList(sdocid)),
        })
    return {"docs": vresult}


def GetDocDetail(sdocid, blight=False):
    sdir = DocDir(sdocid)
    if not os.path.isdir(sdir):
        raise ValueError("文档不存在")
    ometa = ReadJson(os.path.join(sdir, "meta.json"))
    ocomments = ReadJson(os.path.join(sdir, "comments.json"), {"items": []})
    otodos = ReadJson(os.path.join(sdir, "todos.json"), {"items": []})
    vrevs = _LoadRevisionList(sdocid)
    return {
        "meta": ometa,
        "comments": ocomments.get("items", []),
        "todos": otodos.get("items", []),
        "progress": _CalcProgress(otodos),
        "revisions": vrevs[:20],
        "preview_ready": os.path.isfile(os.path.join(sdir, "preview.html")),
        "has_working_stash": HasWorkingStash(sdocid),
        "working_status": GetWorkingStatus(sdocid, bfull=not blight),
    }


def _CalcProgress(otodos):
    vitems = otodos.get("items", [])
    if not vitems:
        return {"done": 0, "total": 0, "percent": 100}
    ndone = sum(1 for x in vitems if x.get("status") == "done")
    return {"done": ndone, "total": len(vitems), "percent": int(ndone * 100 / len(vitems))}


def UpdateDocMeta(sdocid, stitle=None, vtags=None):
    spath = os.path.join(DocDir(sdocid), "meta.json")
    if not os.path.isfile(spath):
        raise ValueError("文档不存在")
    ometa = ReadJson(spath)
    if stitle:
        ometa["title"] = stitle.strip()
    if vtags is not None:
        ometa["tags"] = [t.strip() for t in vtags if t.strip()]
    ometa["updated"] = datetime.now().strftime("%Y-%m-%d %H:%M")
    WriteJson(spath, ometa)
    omanifest = ReadManifest()
    for oitem in omanifest.get("docs", []):
        if oitem["id"] == sdocid:
            oitem["title"] = ometa["title"]
            oitem["updated"] = ometa["updated"]
    WriteManifest(omanifest)
    return ometa


def _TextFromXml(onode):
    vparts = []
    for telem in onode.iter("{%s}t" % WNS):
        if telem.text:
            vparts.append(telem.text)
    return "".join(vparts).strip()


def _ParseCommentsXml(scomments_xml):
    vcomments = []
    if not scomments_xml:
        return vcomments
    oroot = ET.fromstring(scomments_xml)
    for ocomment in oroot.findall("w:comment", NS):
        scid = ocomment.get("{%s}id" % WNS, "")
        sauthor = ocomment.get("{%s}author" % WNS, "")
        sdate = ocomment.get("{%s}date" % WNS, "")
        stext = _TextFromXml(ocomment)
        vcomments.append({
            "id": scid,
            "author": sauthor,
            "date": sdate[:19] if sdate else "",
            "text": stext,
        })
    return vcomments


def _MapCommentParas(sdocument_xml):
    """批注 id → 段落索引。"""
    if not sdocument_xml:
        return {}
    oroot = ET.fromstring(sdocument_xml)
    obody = oroot.find("w:body", NS)
    if obody is None:
        return {}
    omap = {}
    npara = -1
    for ochild in obody:
        stag = ochild.tag.split("}")[-1] if "}" in ochild.tag else ochild.tag
        if stag == "p":
            npara += 1
        for ostart in ochild.iter("{%s}commentRangeStart" % WNS):
            scid = ostart.get("{%s}id" % WNS)
            if scid is not None:
                omap[scid] = npara
    return omap


def _EscHtml(stext):
    return (stext or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")


def _EscAttr(stext):
    return _EscHtml(stext).replace("'", "&#39;")


def _MimeFromExt(sext):
    return {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".gif": "image/gif",
        ".bmp": "image/bmp",
        ".webp": "image/webp",
        ".emf": "image/emf",
        ".wmf": "image/wmf",
    }.get(sext.lower(), "application/octet-stream")


def _LoadImageRels(scurrent):
    with zipfile.ZipFile(scurrent, "r") as oz:
        srels = "word/_rels/document.xml.rels"
        if srels not in oz.namelist():
            return {}
        oroot = ET.fromstring(oz.read(srels))
        omap = {}
        for orel in oroot:
            if orel.get("Type", "").endswith("/image"):
                omap[orel.get("Id")] = orel.get("Target", "")
        return omap


def _LoadImageUrls(sdocid, scurrent):
    """rId / 文件名 → 媒体接口 URL，避免整页内嵌 base64。"""
    orels = _LoadImageRels(scurrent)
    odata = {}
    for srid, starget in orels.items():
        sfname = os.path.basename(starget.replace("\\", "/"))
        if not sfname:
            continue
        surl = "/api/docs/media?id=%s&file=%s" % (
            urllib.parse.quote(sdocid, safe=""),
            urllib.parse.quote(sfname, safe=""),
        )
        odata[srid] = surl
        odata[sfname] = surl
    return odata


def _ParaHasDrawing(opara):
    for orun in opara.runs:
        if orun._element.findall(".//{%s}drawing" % WNS):
            return True
    return False


def _RgbToHex(orgb):
    if not orgb:
        return ""
    try:
        return "#%02x%02x%02x" % (orgb[0], orgb[1], orgb[2])
    except Exception:
        return ""


def _TwipsToPt(ntwips):
    return round(float(ntwips) / 20.0, 1)


def _ElemRPr(oelem):
    if oelem is None:
        return None
    return oelem.find("{%s}rPr" % WNS)


def _ElemPPr(oelem):
    if oelem is None:
        return None
    return oelem.find("{%s}pPr" % WNS)


def _LoadThemeFonts(scurrent):
    othemes = {}
    if not scurrent or not os.path.isfile(scurrent):
        return othemes
    try:
        with zipfile.ZipFile(scurrent, "r") as oz:
            spath = "word/theme/theme1.xml"
            if spath not in oz.namelist():
                return othemes
            oroot = ET.fromstring(oz.read(spath))
        for sprefix in ("major", "minor"):
            ofont = oroot.find(".//{%s}%sFont" % (ANS, sprefix))
            if ofont is None:
                continue
            for stag, skey in (("latin", "latin"), ("ea", "eastAsia"), ("cs", "cs")):
                oel = ofont.find("{%s}%s" % (ANS, stag))
                if oel is not None:
                    stype = (oel.get("typeface") or "").strip()
                    if stype:
                        othemes["%s_%s" % (sprefix, skey)] = stype
    except Exception:
        pass
    return othemes


def _ResolveThemeFontName(stheme_attr, othemes):
    if not stheme_attr or not othemes:
        return ""
    stheme_attr = stheme_attr.strip()
    smap = {
        "majorEastAsia": ("major", "eastAsia"),
        "minorEastAsia": ("minor", "eastAsia"),
        "majorHAnsi": ("major", "latin"),
        "minorHAnsi": ("minor", "latin"),
        "majorAscii": ("major", "latin"),
        "minorAscii": ("minor", "latin"),
        "majorBidi": ("major", "cs"),
        "minorBidi": ("minor", "cs"),
    }
    okey = smap.get(stheme_attr)
    if not okey:
        return ""
    return othemes.get("%s_%s" % okey) or ""


def _RFontsFromRPr(rpr, othemes=None):
    if rpr is None:
        return ""
    rf = rpr.find("{%s}rFonts" % WNS)
    if rf is None:
        return ""
    for sattr in ("eastAsia", "ascii", "hAnsi", "cs"):
        sval = rf.get("{%s}%s" % (WNS, sattr))
        if sval:
            return sval.strip()
    for sattr, stheme in (
        ("eastAsiaTheme", "majorEastAsia"),
        ("asciiTheme", "majorHAnsi"),
        ("hAnsiTheme", "majorHAnsi"),
        ("cstheme", "majorBidi"),
    ):
        sval = rf.get("{%s}%s" % (WNS, sattr))
        if sval:
            sresolved = _ResolveThemeFontName(sval, othemes)
            if sresolved:
                return sresolved
    return ""


def _SzPtFromRPr(rpr):
    if rpr is None:
        return None
    sz = rpr.find("{%s}sz" % WNS)
    if sz is None:
        return None
    sval = sz.get("{%s}val" % WNS)
    if not sval:
        return None
    try:
        return round(int(sval) / 2.0, 1)
    except Exception:
        return None


def _ColorFromRPr(rpr):
    if rpr is None:
        return ""
    ocolor = rpr.find("{%s}color" % WNS)
    if ocolor is None:
        return ""
    sval = (ocolor.get("{%s}val" % WNS) or "").strip().lower()
    if not sval or sval == "auto":
        return ""
    if sval == "000000":
        return "#000000"
    try:
        return "#%s" % sval[-6:].zfill(6)
    except Exception:
        return ""


def _FlagFromRPr(rpr, stag):
    if rpr is None:
        return False
    oel = rpr.find("{%s}%s" % (WNS, stag))
    if oel is None:
        return False
    sval = oel.get("{%s}val" % WNS)
    return sval is None or sval not in ("0", "false", "off")


def _JcValToCss(sval):
    return {
        "left": "left", "start": "left",
        "center": "center",
        "right": "right", "end": "right",
        "both": "justify", "distribute": "justify", "justify": "justify",
    }.get((sval or "").strip().lower(), "")


def _AlignEnumToCss(nalign):
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        return {
            WD_ALIGN_PARAGRAPH.LEFT: "left",
            WD_ALIGN_PARAGRAPH.CENTER: "center",
            WD_ALIGN_PARAGRAPH.RIGHT: "right",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
        }.get(nalign, "")
    except Exception:
        return ""


def _IndLayoutFromPPr(ppr):
    olayout = {}
    if ppr is None:
        return olayout
    oind = ppr.find("{%s}ind" % WNS)
    if oind is None:
        return olayout

    def _TwipsAttr(sattr):
        sval = oind.get("{%s}%s" % (WNS, sattr))
        if not sval:
            return None
        try:
            return _TwipsToPt(int(sval))
        except Exception:
            return None

    nleft = _TwipsAttr("left")
    nright = _TwipsAttr("right")
    nfirst = _TwipsAttr("firstLine")
    nhang = _TwipsAttr("hanging")
    if nleft is not None:
        olayout["left_pt"] = nleft
    if nright is not None:
        olayout["right_pt"] = nright
    if nhang is not None:
        olayout["hang_pt"] = nhang
        if nleft is None:
            olayout["left_pt"] = nhang
    elif nfirst is not None:
        if nfirst >= 0:
            olayout["first_pt"] = nfirst
        else:
            olayout["hang_pt"] = abs(nfirst)
            olayout["left_pt"] = olayout.get("left_pt", abs(nfirst))
    return olayout


def _SpacingLayoutFromPPr(ppr):
    olayout = {}
    if ppr is None:
        return olayout
    osp = ppr.find("{%s}spacing" % WNS)
    if osp is None:
        return olayout
    for sattr, skey in (("before", "space_before_pt"), ("after", "space_after_pt")):
        sval = osp.get("{%s}%s" % (WNS, sattr))
        if sval:
            try:
                olayout[skey] = _TwipsToPt(int(sval))
            except Exception:
                pass
    sline = osp.get("{%s}line" % WNS)
    if sline:
        try:
            nline = int(sline)
            srule = (osp.get("{%s}lineRule" % WNS) or "auto").lower()
            if srule == "auto":
                olayout["line_height"] = round(nline / 240.0, 2)
            else:
                olayout["line_height_pt"] = _TwipsToPt(nline)
        except Exception:
            pass
    return olayout


def _StylePPrRPr(opara):
    oppr, orpr = None, None
    try:
        if opara.style and opara.style._element is not None:
            oppr = _ElemPPr(opara.style._element)
            orpr = _ElemRPr(opara.style._element)
            if orpr is None and oppr is not None:
                orpr = _ElemRPr(oppr)
    except Exception:
        pass
    return oppr, orpr


def _ParaLayoutDict(opara, othemes=None):
    olayout = {}
    pf = opara.paragraph_format
    if pf.alignment is not None:
        scss = _AlignEnumToCss(pf.alignment)
        if scss:
            olayout["align"] = scss
    if pf.left_indent is not None:
        olayout["left_pt"] = round(pf.left_indent.pt, 1)
    if pf.right_indent is not None:
        olayout["right_pt"] = round(pf.right_indent.pt, 1)
    if pf.first_line_indent is not None:
        npt = round(pf.first_line_indent.pt, 1)
        if npt > 0:
            olayout["first_pt"] = npt
        elif npt < 0:
            olayout["hang_pt"] = abs(npt)
            olayout["left_pt"] = olayout.get("left_pt", abs(npt))
    if pf.space_before is not None:
        olayout["space_before_pt"] = round(pf.space_before.pt, 1)
    if pf.space_after is not None:
        olayout["space_after_pt"] = round(pf.space_after.pt, 1)
    if pf.line_spacing is not None:
        try:
            from docx.enum.text import WD_LINE_SPACING
            if pf.line_spacing_rule == WD_LINE_SPACING.MULTIPLE:
                olayout["line_height"] = round(float(pf.line_spacing), 2)
            else:
                olayout["line_height_pt"] = round(pf.line_spacing.pt, 1)
        except Exception:
            pass

    ppr = _ElemPPr(opara._element)
    if ppr is not None:
        ojc = ppr.find("{%s}jc" % WNS)
        if ojc is not None and "align" not in olayout:
            scss = _JcValToCss(ojc.get("{%s}val" % WNS))
            if scss:
                olayout["align"] = scss
        for skey, sval in _IndLayoutFromPPr(ppr).items():
            olayout.setdefault(skey, sval)
        for skey, sval in _SpacingLayoutFromPPr(ppr).items():
            olayout.setdefault(skey, sval)

    sppr, _ = _StylePPrRPr(opara)
    if sppr is not None:
        if "align" not in olayout:
            ojc = sppr.find("{%s}jc" % WNS)
            if ojc is not None:
                scss = _JcValToCss(ojc.get("{%s}val" % WNS))
                if scss:
                    olayout["align"] = scss
        for skey, sval in _IndLayoutFromPPr(sppr).items():
            olayout.setdefault(skey, sval)
        for skey, sval in _SpacingLayoutFromPPr(sppr).items():
            olayout.setdefault(skey, sval)
    return olayout


def _LayoutDictToCss(olayout):
    vstyles = []
    if olayout.get("align"):
        vstyles.append("text-align:%s" % olayout["align"])
    if olayout.get("left_pt") is not None:
        vstyles.append("margin-left:%gpt" % olayout["left_pt"])
    if olayout.get("right_pt") is not None:
        vstyles.append("margin-right:%gpt" % olayout["right_pt"])
    if olayout.get("first_pt") is not None:
        vstyles.append("text-indent:%gpt" % olayout["first_pt"])
    if olayout.get("hang_pt") is not None:
        nhang = olayout["hang_pt"]
        nleft = olayout.get("left_pt", nhang)
        vstyles.append("padding-left:%gpt" % nleft)
        vstyles.append("text-indent:-%gpt" % nhang)
    if olayout.get("space_before_pt") is not None:
        vstyles.append("margin-top:%gpt" % olayout["space_before_pt"])
    if olayout.get("space_after_pt") is not None:
        vstyles.append("margin-bottom:%gpt" % olayout["space_after_pt"])
    if olayout.get("line_height") is not None:
        vstyles.append("line-height:%s" % olayout["line_height"])
    elif olayout.get("line_height_pt") is not None:
        vstyles.append("line-height:%gpt" % olayout["line_height_pt"])
    return ";".join(vstyles)


def _ParaStyleCss(opara, othemes=None):
    return _LayoutDictToCss(_ParaLayoutDict(opara, othemes))


def _RunFormatDict(orun, opara=None, othemes=None):
    ofmt = {
        "bold": bool(orun.bold),
        "italic": bool(orun.italic),
        "underline": bool(orun.underline),
        "font": (orun.font.name or "").strip(),
        "size_pt": None,
        "color": "",
    }
    if orun.font.size:
        try:
            ofmt["size_pt"] = round(orun.font.size.pt, 1)
        except Exception:
            pass
    if orun.font.color and orun.font.color.rgb:
        ofmt["color"] = _RgbToHex(orun.font.color.rgb)

    rr = _ElemRPr(orun._element)
    if rr is not None:
        if not ofmt["bold"] and _FlagFromRPr(rr, "b"):
            ofmt["bold"] = True
        if not ofmt["italic"] and _FlagFromRPr(rr, "i"):
            ofmt["italic"] = True
        if not ofmt["underline"] and _FlagFromRPr(rr, "u"):
            ofmt["underline"] = True
        if not ofmt["font"]:
            ofmt["font"] = _RFontsFromRPr(rr, othemes)
        if ofmt["size_pt"] is None:
            ofmt["size_pt"] = _SzPtFromRPr(rr)
        if not ofmt["color"]:
            ofmt["color"] = _ColorFromRPr(rr)

    ppr = _ElemPPr(opara._element) if opara is not None else None
    pr = _ElemRPr(ppr) if ppr is not None else None
    if pr is not None:
        if not ofmt["font"]:
            ofmt["font"] = _RFontsFromRPr(pr, othemes)
        if ofmt["size_pt"] is None:
            ofmt["size_pt"] = _SzPtFromRPr(pr)
        if not ofmt["color"]:
            ofmt["color"] = _ColorFromRPr(pr)
        if not ofmt["bold"] and _FlagFromRPr(pr, "b"):
            ofmt["bold"] = True
        if not ofmt["italic"] and _FlagFromRPr(pr, "i"):
            ofmt["italic"] = True
        if not ofmt["underline"] and _FlagFromRPr(pr, "u"):
            ofmt["underline"] = True

    _, sr = _StylePPrRPr(opara) if opara is not None else (None, None)
    if sr is not None:
        if not ofmt["font"]:
            ofmt["font"] = _RFontsFromRPr(sr, othemes)
        if ofmt["size_pt"] is None:
            ofmt["size_pt"] = _SzPtFromRPr(sr)
        if not ofmt["color"]:
            ofmt["color"] = _ColorFromRPr(sr)
        if not ofmt["bold"] and _FlagFromRPr(sr, "b"):
            ofmt["bold"] = True
        if not ofmt["italic"] and _FlagFromRPr(sr, "i"):
            ofmt["italic"] = True
        if not ofmt["underline"] and _FlagFromRPr(sr, "u"):
            ofmt["underline"] = True

    if opara is not None:
        try:
            if opara.style and opara.style.font:
                if not ofmt["font"] and opara.style.font.name:
                    ofmt["font"] = opara.style.font.name.strip()
                if ofmt["size_pt"] is None and opara.style.font.size:
                    ofmt["size_pt"] = round(opara.style.font.size.pt, 1)
        except Exception:
            pass
    return ofmt


def _FormatDictToCss(ofmt):
    vstyles = []
    if ofmt.get("font"):
        vstyles.append("font-family:%s" % ofmt["font"])
    if ofmt.get("size_pt"):
        vstyles.append("font-size:%gpt" % ofmt["size_pt"])
    if ofmt.get("color"):
        vstyles.append("color:%s" % ofmt["color"])
    return ";".join(vstyles)


def _WrapText(stext, orun, opara=None, othemes=None):
    if not stext:
        return ""
    ofmt = _RunFormatDict(orun, opara, othemes)
    shtml = _EscHtml(stext)
    if ofmt.get("bold"):
        shtml = "<strong>%s</strong>" % shtml
    if ofmt.get("italic"):
        shtml = "<em>%s</em>" % shtml
    if ofmt.get("underline"):
        shtml = "<u>%s</u>" % shtml
    sstyle = _FormatDictToCss(ofmt)
    if sstyle:
        shtml = '<span style="%s">%s</span>' % (_EscAttr(sstyle), shtml)
    return shtml


def _RunToHtml(orun, oimgurls, beditable=False, opara=None, othemes=None):
    vparts = []
    if orun.text:
        vparts.append(_WrapText(orun.text, orun, opara, othemes))
    for oblip in orun._element.iter("{%s}blip" % ANS):
        srid = oblip.get("{%s}embed" % RNS)
        surl = oimgurls.get(srid, "")
        if surl:
            sattr = ' contenteditable="false"' if beditable else ""
            vparts.append(
                '<span class="imgwrap"%s><img class="docimg" src="%s" alt=""></span>' % (sattr, surl)
            )
    return "".join(vparts)


def _ParaToHtml(opara, oimgurls, beditable=False, othemes=None):
    if not opara.runs:
        stext = _EscHtml(opara.text or "") or "&#160;"
        sstyle = _ParaStyleCss(opara, othemes)
        if sstyle and stext != "&#160;":
            return '<span style="%s">%s</span>' % (_EscAttr(sstyle), stext)
        return stext
    sbody = "".join(_RunToHtml(r, oimgurls, beditable, opara, othemes) for r in opara.runs)
    return sbody or "&#160;"


def _TableToHtml(otable, oimgurls, beditable=False, othemes=None):
    vrows = []
    for orow in otable.rows:
        vcells = []
        for ocell in orow.cells:
            vparts = []
            for opara in ocell.paragraphs:
                vparts.append(_ParaToHtml(opara, oimgurls, beditable, othemes))
            vcells.append("<td>%s</td>" % ("<br>".join(vparts) or "&#160;"))
        vrows.append("<tr>%s</tr>" % "".join(vcells))
    return '<table class="doctable">%s</table>' % "".join(vrows)


def _CommentMarks(npara, opara_comments):
    if npara not in opara_comments:
        return ""
    return "".join(
        '<span class="cmtmark" data-cid="%s" title="%s">📝</span>' % (
            c["id"],
            (c["text"][:40] + "…").replace('"', "'"),
        )
        for c in opara_comments[npara]
    )


def _SanitizeParaText(stext):
    """去除预览 UI 泄漏的批注图标等，不写入 docx。"""
    if not stext:
        return ""
    return stext.replace("\U0001f4dd", "").replace("📝", "").strip()


_FONT_SIZE_LEGACY = {1: 8, 2: 10, 3: 12, 4: 14, 5: 18, 6: 24, 7: 36}


def _ParseParaCss(sstyle):
    olayout = {}
    if not sstyle:
        return olayout
    for spart in sstyle.split(";"):
        if ":" not in spart:
            continue
        skey, sval = spart.split(":", 1)
        skey = skey.strip().lower()
        sval = sval.strip().strip("'\"")
        if skey == "text-align":
            olayout["align"] = sval
        elif skey == "margin-left":
            npt = _ParseFontSize(sval)
            if npt is not None:
                olayout["left_pt"] = npt
        elif skey == "margin-right":
            npt = _ParseFontSize(sval)
            if npt is not None:
                olayout["right_pt"] = npt
        elif skey == "text-indent":
            npt = _ParseFontSize(sval)
            if npt is not None:
                if npt < 0:
                    olayout["hang_pt"] = abs(npt)
                else:
                    olayout["first_pt"] = npt
        elif skey == "padding-left":
            npt = _ParseFontSize(sval)
            if npt is not None:
                olayout["left_pt"] = npt
        elif skey == "margin-top":
            npt = _ParseFontSize(sval)
            if npt is not None:
                olayout["space_before_pt"] = npt
        elif skey == "margin-bottom":
            npt = _ParseFontSize(sval)
            if npt is not None:
                olayout["space_after_pt"] = npt
        elif skey == "line-height":
            try:
                if sval.endswith("pt"):
                    olayout["line_height_pt"] = _ParseFontSize(sval)
                else:
                    olayout["line_height"] = round(float(sval), 2)
            except Exception:
                pass
    return olayout


def _ApplyParaLayout(opara, scss):
    olayout = _ParseParaCss(scss)
    if not olayout:
        return
    pf = opara.paragraph_format
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    from docx.enum.text import WD_LINE_SPACING
    omap = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    if "align" in olayout:
        pf.alignment = omap.get(olayout["align"], WD_ALIGN_PARAGRAPH.LEFT)
    if "left_pt" in olayout:
        pf.left_indent = Pt(olayout["left_pt"])
    if "right_pt" in olayout:
        pf.right_indent = Pt(olayout["right_pt"])
    if "hang_pt" in olayout:
        pf.first_line_indent = Pt(-olayout["hang_pt"])
    elif "first_pt" in olayout:
        pf.first_line_indent = Pt(olayout["first_pt"])
    if "space_before_pt" in olayout:
        pf.space_before = Pt(olayout["space_before_pt"])
    if "space_after_pt" in olayout:
        pf.space_after = Pt(olayout["space_after_pt"])
    if "line_height" in olayout:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = olayout["line_height"]
    elif "line_height_pt" in olayout:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(olayout["line_height_pt"])


def _ParseCssStyle(sstyle):
    ostyle = {}
    if not sstyle:
        return ostyle
    for spart in sstyle.split(";"):
        if ":" not in spart:
            continue
        skey, sval = spart.split(":", 1)
        skey = skey.strip().lower()
        sval = sval.strip().strip("'\"")
        if skey == "font-family":
            sval = sval.split(",")[0].strip()
            if sval:
                ostyle["font"] = sval
        elif skey == "font-size":
            osize = _ParseFontSize(sval)
            if osize:
                ostyle["size_pt"] = osize
        elif skey == "color":
            orgb = _ParseColor(sval)
            if orgb:
                ostyle["color"] = orgb
        elif skey == "font-weight" and sval in ("bold", "700", "800", "900"):
            ostyle["bold"] = True
        elif skey == "font-style" and sval == "italic":
            ostyle["italic"] = True
        elif skey == "text-decoration" and "underline" in sval:
            ostyle["underline"] = True
    return ostyle


def _ParseFontSize(sval):
    if not sval:
        return None
    sm = re.match(r"^([\d.]+)\s*(pt|px)?$", sval.lower())
    if not sm:
        return None
    nval = float(sm.group(1))
    sunit = sm.group(2) or "pt"
    if sunit == "px":
        nval = nval * 0.75
    return round(nval, 1) if nval > 0 else None


def _ParseColor(sval):
    if not sval:
        return None
    sval = sval.strip().lower()
    sm = re.match(r"^#([0-9a-f]{6})$", sval)
    if sm:
        shex = sm.group(1)
        return (int(shex[0:2], 16), int(shex[2:4], 16), int(shex[4:6], 16))
    sm = re.match(r"^rgb\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)$", sval)
    if sm:
        return (int(sm.group(1)), int(sm.group(2)), int(sm.group(3)))
    omap = {
        "black": (0, 0, 0), "white": (255, 255, 255), "red": (255, 0, 0),
        "blue": (0, 0, 255), "green": (0, 128, 0), "gray": (128, 128, 128),
        "grey": (128, 128, 128),
    }
    return omap.get(sval)


def _MergeStyle(obase, oextra):
    oout = dict(obase or {})
    for skey, sval in (oextra or {}).items():
        if sval is not None:
            oout[skey] = sval
    return oout


class _InlineHtmlParser(HTMLParser):
    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.vsegments = []
        self.vstack = [{}]

    def _PushStyle(self, oextra):
        self.vstack.append(_MergeStyle(self.vstack[-1], oextra))

    def handle_starttag(self, stag, vattrs):
        stag = stag.lower()
        oattrs = dict(vattrs)
        if stag in ("b", "strong"):
            self._PushStyle({"bold": True})
        elif stag in ("i", "em"):
            self._PushStyle({"italic": True})
        elif stag == "u":
            self._PushStyle({"underline": True})
        elif stag == "span":
            self._PushStyle(_ParseCssStyle(oattrs.get("style", "")))
        elif stag == "font":
            oextra = {}
            if oattrs.get("face"):
                oextra["font"] = oattrs.get("face").split(",")[0].strip()
            if oattrs.get("color"):
                orgb = _ParseColor(oattrs.get("color"))
                if orgb:
                    oextra["color"] = orgb
            if oattrs.get("size"):
                try:
                    nlegacy = _FONT_SIZE_LEGACY.get(int(oattrs.get("size")))
                    if nlegacy:
                        oextra["size_pt"] = float(nlegacy)
                except Exception:
                    pass
            self._PushStyle(oextra)
        elif stag == "br":
            self.vsegments.append(_MergeStyle(self.vstack[-1], {"text": "\n"}))

    def handle_endtag(self, stag):
        stag = stag.lower()
        if stag in ("b", "strong", "i", "em", "u", "span", "font") and len(self.vstack) > 1:
            self.vstack.pop()

    def handle_data(self, sdata):
        if not sdata:
            return
        self.vsegments.append(_MergeStyle(self.vstack[-1], {"text": sdata}))


def _NormalizeEditorHtml(shtml):
    if not shtml:
        return ""
    shtml = shtml.strip()
    shtml = re.sub(r"</?(?:div|p)\b[^>]*>", "", shtml, flags=re.I)
    shtml = re.sub(r"<br\s*/?>", "<br>", shtml, flags=re.I)
    return shtml.strip()


def _ParseInlineHtml(shtml):
    sparser = _InlineHtmlParser()
    try:
        sparser.feed(_NormalizeEditorHtml(shtml))
        sparser.close()
    except Exception:
        return [{"text": _SanitizeParaText(re.sub(r"<[^>]+>", "", shtml or ""))}]
    vmerged = []
    for oseg in sparser.vsegments:
        stext = oseg.get("text", "")
        if not stext:
            continue
        ostyle = {k: v for k, v in oseg.items() if k != "text"}
        if vmerged:
            oprev = {k: v for k, v in vmerged[-1].items() if k != "text"}
            if oprev == ostyle:
                vmerged[-1]["text"] += stext
                continue
        vmerged.append(dict(ostyle, text=stext))
    return vmerged or [{"text": ""}]


def _ClearParaRuns(opara):
    for orun in list(opara.runs):
        orun._element.getparent().remove(orun._element)


def _ApplySegmentStyle(orun, oseg):
    if oseg.get("bold"):
        orun.bold = True
    if oseg.get("italic"):
        orun.italic = True
    if oseg.get("underline"):
        orun.underline = True
    if oseg.get("font"):
        orun.font.name = oseg["font"]
    if oseg.get("size_pt"):
        from docx.shared import Pt
        orun.font.size = Pt(oseg["size_pt"])
    if oseg.get("color"):
        from docx.shared import RGBColor
        orgb = oseg["color"]
        orun.font.color.rgb = RGBColor(orgb[0], orgb[1], orgb[2])


def _ApplyRichHtmlToPara(opara, shtml):
    vsegments = _ParseInlineHtml(shtml)
    if _ParaHasDrawing(opara):
        stext = _SanitizeParaText("".join(s.get("text", "") for s in vsegments))
        _SetParaTextPreserveMedia(opara, stext)
        return
    _ClearParaRuns(opara)
    odoc = opara.part.document if hasattr(opara, "part") else None
    for oseg in vsegments:
        stext = oseg.get("text", "")
        if not stext:
            continue
        orun = opara.add_run(stext)
        _ApplySegmentStyle(orun, oseg)


def _ParaComparable(opara, othemes=None):
    vruns = []
    for orun in opara.runs:
        ofmt = _RunFormatDict(orun, opara, othemes)
        vruns.append({
            "t": orun.text or "",
            "b": ofmt.get("bold"),
            "i": ofmt.get("italic"),
            "u": ofmt.get("underline"),
            "f": ofmt.get("font") or "",
            "s": ofmt.get("size_pt") or "",
            "c": ofmt.get("color") or "",
        })
    return json.dumps({
        "layout": _ParaLayoutDict(opara, othemes),
        "runs": vruns,
    }, ensure_ascii=False)


def _IterAllParagraphs(odoc):
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    for ochild in odoc.element.body:
        stag = ochild.tag.split("}")[-1] if "}" in ochild.tag else ochild.tag
        if stag == "p":
            yield Paragraph(ochild, odoc)
        elif stag == "tbl":
            otable = Table(ochild, odoc)
            for orow in otable.rows:
                for ocell in orow.cells:
                    for opara in ocell.paragraphs:
                        yield opara


def _CleanDocxUiArtifacts(odoc):
    bchanged = False
    for opara in _IterAllParagraphs(odoc):
        stext = opara.text or ""
        sclean = _SanitizeParaText(stext)
        if sclean == stext:
            continue
        bchanged = True
        if _ParaHasDrawing(opara):
            _SetParaTextPreserveMedia(opara, sclean)
        else:
            opara.text = sclean
    return bchanged


def _ParaBlockHtml(npara, opara, oimgurls, opara_comments_all, opara_comments_pending, beditable=False, othemes=None):
    sbody = _ParaToHtml(opara, oimgurls, beditable, othemes)
    scls = "docpara"
    if beditable:
        scls += " docpara-editable"
    if npara in opara_comments_pending:
        scls += " has-comment"
    if _ParaHasDrawing(opara):
        scls += " has-image"
    sedit = ' contenteditable="true"' if beditable else ""
    splain = _EscAttr(_SanitizeParaText(opara.text or ""))
    spstyle = _ParaStyleCss(opara, othemes)
    sstyleattr = (' style="%s"' % _EscAttr(spstyle)) if spstyle else ""
    spstyleattr = (' data-pstyle="%s"' % _EscAttr(spstyle)) if spstyle else ""
    smarks = _CommentMarks(npara, opara_comments_all)
    if beditable and smarks:
        return (
            '<div class="parablock" data-para="%d">'
            '<div class="cmtmarks" contenteditable="false">%s</div>'
            '<p id="para-%d" class="%s" data-para="%d" data-plain="%s"%s%s%s>%s</p>'
            '</div>'
        ) % (npara, smarks, npara, scls, npara, splain, spstyleattr, sstyleattr, sedit, sbody)
    return (
        '<p id="para-%d" class="%s" data-para="%d" data-plain="%s"%s%s%s>%s%s</p>'
        % (npara, scls, npara, splain, spstyleattr, sstyleattr, sedit, smarks, sbody)
    )


def _BuildDocumentBodyHtml(odoc, oimgurls, opara_comments_all, opara_comments_pending, beditable=False, othemes=None):
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    vhtml = []
    npara = -1
    for ochild in odoc.element.body:
        stag = ochild.tag.split("}")[-1] if "}" in ochild.tag else ochild.tag
        if stag == "p":
            npara += 1
            opara = Paragraph(ochild, odoc)
            vhtml.append(_ParaBlockHtml(
                npara, opara, oimgurls, opara_comments_all, opara_comments_pending, beditable, othemes,
            ))
        elif stag == "tbl":
            vhtml.append(_TableToHtml(Table(ochild, odoc), oimgurls, beditable, othemes))
    return "\n".join(vhtml)


def _SetParaTextPreserveMedia(opara, stext):
    vtext_runs = []
    for orun in opara.runs:
        if not orun._element.findall(".//{%s}drawing" % WNS):
            vtext_runs.append(orun)
    if not vtext_runs:
        if stext:
            opara.add_run(stext)
        return
    vtext_runs[0].text = stext
    for orun in vtext_runs[1:]:
        orun._element.getparent().remove(orun._element)


def GetMediaBytes(sdocid, sfname):
    sfname = os.path.basename((sfname or "").replace("\\", "/"))
    if not sfname or ".." in sfname:
        raise ValueError("无效的图片名")
    scurrent = os.path.join(DocDir(sdocid), "current.docx")
    if not os.path.isfile(scurrent):
        raise ValueError("文档不存在")
    with zipfile.ZipFile(scurrent, "r") as oz:
        for spath in ("word/media/%s" % sfname, "word/%s" % sfname):
            if spath in oz.namelist():
                bdata = oz.read(spath)
                return bdata, _MimeFromExt(os.path.splitext(sfname)[1])
    raise ValueError("图片不存在")


def _ParaExcerpt(sdocx_path, nindex):
    from docx import Document
    odoc = Document(sdocx_path)
    if nindex < 0 or nindex >= len(odoc.paragraphs):
        return ""
    return (odoc.paragraphs[nindex].text or "")[:120]


def ExtractComments(sdocid):
    sdir = DocDir(sdocid)
    scurrent = os.path.join(sdir, "current.docx")
    if not os.path.isfile(scurrent):
        raise ValueError("文档文件不存在")
    with zipfile.ZipFile(scurrent, "r") as oz:
        scomments = oz.read("word/comments.xml").decode("utf-8") if "word/comments.xml" in oz.namelist() else ""
        sdocument = oz.read("word/document.xml").decode("utf-8")
    vraw = _ParseCommentsXml(scomments)
    oparas = _MapCommentParas(sdocument)
    vitems = []
    for i, oc in enumerate(vraw):
        scid = oc["id"]
        npara = oparas.get(scid, -1)
        vitems.append({
            "id": scid or str(i),
            "author": oc["author"],
            "date": oc["date"],
            "text": oc["text"],
            "para_index": npara,
            "excerpt": _ParaExcerpt(scurrent, npara) if npara >= 0 else "",
            "status": "pending",
        })
    otodos_old = ReadJson(os.path.join(sdir, "todos.json"), {"items": []})
    odone_map = {x.get("comment_id"): x.get("status") for x in otodos_old.get("items", []) if x.get("status") == "done"}
    for oc in vitems:
        if oc["id"] in odone_map:
            oc["status"] = "done"
    WriteJson(os.path.join(sdir, "comments.json"), {
        "extracted": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "items": vitems,
    })
    vtodos = []
    for oc in vitems:
        vtodos.append({
            "id": "todo-%s" % oc["id"],
            "comment_id": oc["id"],
            "text": "[%s] %s" % (oc["author"] or "批注", oc["text"]),
            "para_index": oc["para_index"],
            "status": oc["status"],
        })
    WriteJson(os.path.join(sdir, "todos.json"), {
        "updated": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "items": vtodos,
    })
    BuildPreview(sdocid, vitems)
    ometa = ReadJson(os.path.join(sdir, "meta.json"))
    ometa["updated"] = datetime.now().strftime("%Y-%m-%d %H:%M")
    ometa["comment_count"] = len(vitems)
    WriteJson(os.path.join(sdir, "meta.json"), ometa)
    return {"comments": len(vitems), "todos": len(vtodos)}


def _PparaCommentsMap(vcomments, sstatus=None):
    opara_comments = {}
    for oc in vcomments or []:
        if sstatus and oc.get("status") != sstatus:
            continue
        nidx = oc.get("para_index", -1)
        if nidx >= 0:
            opara_comments.setdefault(nidx, []).append(oc)
    return opara_comments


def BuildPreview(sdocid, vcomments=None):
    from docx import Document
    sdir = DocDir(sdocid)
    scurrent = os.path.join(sdir, "current.docx")
    if vcomments is None:
        vcomments = ReadJson(os.path.join(sdir, "comments.json"), {"items": []}).get("items", [])
    odoc = Document(scurrent)
    oimgurls = _LoadImageUrls(sdocid, scurrent)
    othemes = _LoadThemeFonts(scurrent)
    sbody = _BuildDocumentBodyHtml(
        odoc, oimgurls,
        _PparaCommentsMap(vcomments),
        _PparaCommentsMap(vcomments, "pending"),
        False, othemes,
    )
    shtml = '<div class="docpreview">%s</div>' % sbody
    with open(os.path.join(sdir, "preview.html"), "w", encoding="utf-8") as f:
        f.write(shtml)
    return shtml


_EDITOR_THEME_CSS = {
    "fresh": {
        "workspace": "#e6ebe6", "page": "#f8f9f7", "text": "#3f4a44",
        "focus_bg": "rgba(122,148,136,.1)", "focus_border": "#7a9488",
        "comment_bg": "rgba(184,168,120,.14)", "comment_border": "#b8a878",
        "status_bg": "rgba(248,249,247,.94)", "status_border": "rgba(122,148,136,.22)",
        "shadow": "rgba(63,74,68,.08)",
    },
    "girly": {
        "workspace": "#e8e2de", "page": "#fffcfb", "text": "#4a3f47",
        "focus_bg": "rgba(201,120,154,.1)", "focus_border": "#c9789a",
        "comment_bg": "rgba(232,184,109,.18)", "comment_border": "#e8b86d",
        "status_bg": "rgba(255,252,251,.92)", "status_border": "rgba(201,120,154,.25)",
        "shadow": "rgba(74,63,71,.12)",
    },
    "boyish": {
        "workspace": "#d0d8e4", "page": "#ffffff", "text": "#1e3a5f",
        "focus_bg": "rgba(61,125,214,.1)", "focus_border": "#3d7dd6",
        "comment_bg": "rgba(232,160,64,.16)", "comment_border": "#e8a040",
        "status_bg": "rgba(255,255,255,.94)", "status_border": "rgba(61,125,214,.28)",
        "shadow": "rgba(30,58,95,.12)",
    },
    "cool": {
        "workspace": "#2a3038", "page": "#1c2330", "text": "#e6edf3",
        "focus_bg": "rgba(0,212,255,.12)", "focus_border": "#00d4ff",
        "comment_bg": "rgba(251,191,36,.14)", "comment_border": "#fbbf24",
        "status_bg": "rgba(22,27,34,.92)", "status_border": "rgba(0,212,255,.28)",
        "shadow": "rgba(0,0,0,.35)",
    },
}


def GetEditorCss(stheme="girly"):
    otheme = _EDITOR_THEME_CSS.get(stheme) or _EDITOR_THEME_CSS["girly"]
    return (
        "*{box-sizing:border-box}"
        "html,body{margin:0;height:100%%;overflow:hidden;background:%(workspace)s;font-family:"
        '"PingFang SC","Microsoft YaHei","SimSun",serif;display:flex;flex-direction:column}'
        ".fmtshell{position:fixed;top:0;left:0;right:0;z-index:30;transition:transform .2s ease;will-change:transform}"
        ".fmtshell.fmt-hidden{transform:translateY(-100%%)}"
        ".fmtshell.fmt-hidden.fmt-open,.fmtshell.fmt-hidden:hover{transform:translateY(0)}"
        ".fmtpeek{position:fixed;top:0;left:0;right:0;height:10px;z-index:29;display:none}"
        ".fmtpeek.show{display:block}"
        ".fmtbar{display:flex;align-items:center;gap:6px;flex-wrap:wrap;"
        "padding:6px 12px;background:%(status_bg)s;border-bottom:1px solid %(status_border)s;"
        "box-shadow:0 4px 16px %(shadow)s;backdrop-filter:blur(8px)}"
        ".fmtbar .grp{display:flex;align-items:center;gap:4px;padding-right:8px;margin-right:4px;"
        "border-right:1px solid %(status_border)s}"
        ".fmtbar .grp:last-child{border-right:none;margin-right:0;padding-right:0}"
        ".fmtbtn{min-width:30px;height:30px;padding:0 8px;border-radius:8px;border:1px solid %(status_border)s;"
        "background:%(page)s;color:%(text)s;font-size:13px;font-weight:700;cursor:pointer;line-height:1}"
        ".fmtbtn:hover{border-color:%(focus_border)s;color:%(focus_border)s}"
        ".fmtbtn.on{background:%(focus_bg)s;border-color:%(focus_border)s;color:%(focus_border)s}"
        ".fmtbtn.i{font-style:italic;font-family:Georgia,serif}"
        ".fmtbtn.u{text-decoration:underline}"
        ".fmtselect{height:30px;padding:0 8px;border-radius:8px;border:1px solid %(status_border)s;"
        "background:%(page)s;color:%(text)s;font-size:12px;max-width:118px}"
        ".fmtcolor{width:30px;height:30px;padding:2px;border-radius:8px;border:1px solid %(status_border)s;"
        "background:%(page)s;cursor:pointer}"
        ".docworkspace{flex:1;min-height:0;overflow-y:auto;padding:44px 16px 48px}"
        ".docpage{max-width:794px;margin:0 auto;background:%(page)s;color:%(text)s;"
        "min-height:1123px;padding:72px 84px;box-shadow:0 8px 32px %(shadow)s;"
        "border-radius:4px;font-size:12pt;line-height:1.5}"
        ".parablock{margin:0 0 10px}"
        ".docpage .docpara{font-size:inherit;line-height:inherit}"
        ".cmtmarks{line-height:1.2;margin-bottom:2px;user-select:none}"
        ".docpara{margin:0;padding:2px 4px;border-radius:6px;outline:none;position:relative;min-height:1.4em;"
        "-webkit-user-select:text;user-select:text}"
        ".docpara:focus,.docpara.selected{background:%(focus_bg)s;box-shadow:inset 0 0 0 2px %(focus_border)s}"
        ".docpara.has-comment{background:%(comment_bg)s;border-left:3px solid %(comment_border)s;"
        "padding-left:8px;margin-left:-4px}"
        ".docpara.has-image{min-height:24px}"
        ".imgwrap{display:inline-block;max-width:100%%;vertical-align:middle}"
        ".docimg{max-width:100%%;height:auto;display:block;margin:10px auto;border-radius:2px}"
        ".doctable{width:100%%;border-collapse:collapse;margin:12px 0;font-size:14px}"
        ".doctable td,.doctable th{border:1px solid #ccc;padding:6px 8px;vertical-align:top}"
        ".cmtmark{cursor:pointer;margin-right:6px;user-select:none}"
        ".docstatus{position:fixed;bottom:12px;right:16px;font-size:12px;color:%(text)s;"
        "background:%(status_bg)s;border:1px solid %(status_border)s;padding:6px 14px;"
        "border-radius:999px;box-shadow:0 4px 16px %(shadow)s;display:none}"
        ".docstatus.show{display:block}"
    ) % otheme

_EDITOR_JS = """
const DOC_ID=%(docid)s;
let selectedPara=-1,selectedComment=null,saving=false,activePara=null,savedRange=null;
function notify(o){try{parent.postMessage(Object.assign({source:'paper-doc-editor'},o),'*')}catch(e){}}
function paraFromNode(n){
  while(n){if(n.nodeType===1&&n.classList&&n.classList.contains('docpara-editable'))return n;n=n.parentNode}
  return null;
}
function touchPara(el,scroll){
  if(!el)return;
  document.querySelectorAll('.docpara.selected').forEach(p=>p.classList.remove('selected'));
  el.classList.add('selected');
  activePara=el;
  selectedPara=parseInt(el.dataset.para,10);
  if(scroll!==false)el.scrollIntoView({behavior:'smooth',block:'center'});
}
function selectPara(el,scroll){
  touchPara(el,scroll);
  notify({type:'doc-para',para:selectedPara,plain:el.dataset.plain||'',hasImage:el.classList.contains('has-image')});
  updateFmtUi();
}
function saveSelection(){
  const osel=window.getSelection();
  if(!osel||!osel.rangeCount)return;
  const orange=osel.getRangeAt(0);
  const opara=paraFromNode(orange.commonAncestorContainer);
  if(!opara)return;
  activePara=opara;
  touchPara(opara,false);
  if(!orange.collapsed)savedRange=orange.cloneRange();
}
function restoreSavedSelection(){
  if(!savedRange||!activePara||!document.contains(activePara))return false;
  try{
    activePara.focus({preventScroll:true});
    const osel=window.getSelection();
    osel.removeAllRanges();
    osel.addRange(savedRange);
    return !savedRange.collapsed;
  }catch(e){savedRange=null;return false}
}
function currentSelectionRange(){
  const osel=window.getSelection();
  if(!osel||!osel.rangeCount)return null;
  const orange=osel.getRangeAt(0);
  if(orange.collapsed)return savedRange&&!savedRange.collapsed?savedRange:null;
  return orange;
}
function ensureFmtReady(){
  if(restoreSavedSelection())return activePara;
  if(activePara&&document.contains(activePara)){
    activePara.focus({preventScroll:true});
    return activePara;
  }
  const op=document.querySelector('.docpara-editable.selected')||document.querySelector('.docpara-editable');
  if(op){touchPara(op,false);op.focus({preventScroll:true});return op}
  return null;
}
function needSelection(){
  const orange=currentSelectionRange();
  if(orange&&!orange.collapsed)return orange;
  showStatus('请先框选要排版的文字');
  return null;
}
function wrapRangeStyle(orange,sstyle){
  if(!orange||orange.collapsed)return false;
  const ospan=document.createElement('span');
  ospan.setAttribute('style',sstyle);
  try{orange.surroundContents(ospan)}catch(e){
    const sfrag=orange.extractContents();
    ospan.appendChild(sfrag);
    orange.insertNode(ospan);
  }
  const osel=window.getSelection();
  osel.removeAllRanges();
  const nr=document.createRange();
  nr.selectNodeContents(ospan);
  osel.addRange(nr);
  savedRange=nr.cloneRange();
  return true;
}
function paraPlainText(el){
  const s=(el.innerText||'').replace(/\\uFE0F/g,'').replace(/📝/g,'').replace(/\\s+$/,'');
  return s;
}
function paraHtml(el){
  const oclone=el.cloneNode(true);
  oclone.querySelectorAll('.cmtmark,.imgwrap').forEach(n=>n.remove());
  return oclone.innerHTML.trim();
}
function paraBlockStyle(el){
  return (el.getAttribute('style')||el.dataset.pstyle||'').trim();
}
function initParaState(el){
  if(!el.dataset.html)el.dataset.html=paraHtml(el);
  if(!el.dataset.plain)el.dataset.plain=paraPlainText(el);
  if(!el.dataset.pstyle)el.dataset.pstyle=paraBlockStyle(el);
}
async function savePara(el){
  if(saving)return;
  const n=parseInt(el.dataset.para,10);
  const stext=paraPlainText(el);
  const shtml=paraHtml(el);
  const spstyle=paraBlockStyle(el);
  if(stext===(el.dataset.plain||'')&&shtml===(el.dataset.html||'')&&spstyle===(el.dataset.pstyle||''))return;
  saving=true;showStatus('正在保存…');
  try{
    const obody={id:DOC_ID,para_index:n,text:stext,html:shtml,para_style:spstyle};
    if(selectedComment)obody.comment_id=selectedComment;
    await parent.Api('/api/docs/edit',obody);
    el.dataset.plain=stext;
    el.dataset.html=shtml;
    el.dataset.pstyle=spstyle;
    notify({type:'doc-saved',para:n,commentId:selectedComment||null});
    showStatus('已保存');
  }catch(e){showStatus('保存失败');notify({type:'doc-error',msg:e.message});}
  saving=false;
}
function showStatus(s){const el=document.getElementById('docstatus');el.textContent=s;el.classList.add('show');setTimeout(()=>el.classList.remove('show'),1800);}
function runFmt(cmd,val){
  if(!ensureFmtReady())return;
  const orange=needSelection();if(!orange)return;
  const osel=window.getSelection();
  osel.removeAllRanges();
  osel.addRange(orange);
  try{document.execCommand('styleWithCSS',false,true)}catch(e){}
  document.execCommand(cmd,false,val||null);
  saveSelection();
  updateFmtUi();
}
function applyFontName(sname){
  if(!sname)return;
  if(!ensureFmtReady())return;
  const orange=needSelection();if(!orange)return;
  const osel=window.getSelection();
  osel.removeAllRanges();
  osel.addRange(orange);
  wrapRangeStyle(orange,"font-family:'"+sname.replace(/'/g,"")+"'");
  saveSelection();
  updateFmtUi();
}
function applyFontSize(spt){
  if(!spt)return;
  if(!ensureFmtReady())return;
  const orange=needSelection();if(!orange)return;
  const osel=window.getSelection();
  osel.removeAllRanges();
  osel.addRange(orange);
  wrapRangeStyle(orange,'font-size:'+spt+'pt');
  saveSelection();
  updateFmtUi();
}
function applyColor(scolor){
  if(!scolor)return;
  if(!ensureFmtReady())return;
  const orange=needSelection();if(!orange)return;
  const osel=window.getSelection();
  osel.removeAllRanges();
  osel.addRange(orange);
  try{document.execCommand('styleWithCSS',false,true)}catch(e){}
  if(!document.execCommand('foreColor',false,scolor))wrapRangeStyle(orange,'color:'+scolor);
  saveSelection();
  updateFmtUi();
}
function updateFmtUi(){
  const ob=document.getElementById('fmt_bold');
  const oi=document.getElementById('fmt_italic');
  const ou=document.getElementById('fmt_underline');
  let b=false,i=false,u=false;
  try{
    b=document.queryCommandState('bold');
    i=document.queryCommandState('italic');
    u=document.queryCommandState('underline');
  }catch(e){}
  if(ob)ob.classList.toggle('on',b);
  if(oi)oi.classList.toggle('on',i);
  if(ou)ou.classList.toggle('on',u);
}
function InitFmtReveal(){
  const ows=document.querySelector('.docworkspace');
  const osh=document.querySelector('.fmtshell');
  const opk=document.querySelector('.fmtpeek');
  if(!ows||!osh)return;
  let nlast=0,ntimer=0;
  function SetFmtOpen(b){osh.classList.toggle('fmt-open',!!b)}
  function SetFmtHidden(b){
    osh.classList.toggle('fmt-hidden',!!b);
    if(opk)opk.classList.toggle('show',!!b);
    if(!b)SetFmtOpen(false);
  }
  function KeepFmtOpen(){
    if(osh.classList.contains('fmt-hidden'))SetFmtOpen(true);
  }
  ows.addEventListener('scroll',()=>{
    const nst=ows.scrollTop;
    if(nst<=4){SetFmtHidden(false);SetFmtOpen(false);nlast=nst;return}
    if(nst>nlast+5)SetFmtHidden(true);
    else if(nst<nlast-5){SetFmtHidden(true);SetFmtOpen(true)}
    nlast=nst;
  },{passive:true});
  if(opk){
    opk.addEventListener('mouseenter',KeepFmtOpen);
    opk.addEventListener('click',KeepFmtOpen);
  }
  osh.addEventListener('mouseenter',KeepFmtOpen);
  osh.addEventListener('mouseleave',()=>{
    if(osh.querySelector('select:focus,input:focus'))return;
    clearTimeout(ntimer);
    ntimer=setTimeout(()=>{if(ows.scrollTop>4&&!osh.matches(':hover')&&!opk.matches(':hover'))SetFmtOpen(false)},400);
  });
  osh.addEventListener('focusin',KeepFmtOpen);
}
function bindFmtBar(){
  const osh=document.querySelector('.fmtshell');
  const obar=document.querySelector('.fmtbar');
  if(obar){
    obar.addEventListener('mousedown',e=>{
      const stag=(e.target.tagName||'').toUpperCase();
      if(stag==='SELECT'||stag==='INPUT'||stag==='OPTION')return;
      e.preventDefault();
    });
  }
  InitFmtReveal();
  document.getElementById('fmt_bold').onclick=()=>runFmt('bold');
  document.getElementById('fmt_italic').onclick=()=>runFmt('italic');
  document.getElementById('fmt_underline').onclick=()=>runFmt('underline');
  const ofont=document.getElementById('fmt_font');
  const osize=document.getElementById('fmt_size');
  const ocolor=document.getElementById('fmt_color');
  [ofont,osize,ocolor].forEach(oel=>{
    if(!oel)return;
    oel.addEventListener('mousedown',saveSelection);
    oel.addEventListener('focus',()=>{if(osh)osh.classList.add('fmt-open')});
  });
  ofont.addEventListener('change',e=>{const v=e.target.value;if(v)applyFontName(v);e.target.selectedIndex=0});
  osize.addEventListener('change',e=>{const v=e.target.value;if(v)applyFontSize(v);e.target.selectedIndex=0});
  ocolor.addEventListener('input',e=>applyColor(e.target.value));
  document.addEventListener('selectionchange',()=>{saveSelection();updateFmtUi()});
}
window.focusPara=function(npara,scid){
  selectedComment=scid||null;
  const el=document.getElementById('para-'+npara);
  if(el)selectPara(el);
};
document.querySelectorAll('.docpara-editable').forEach(el=>{
  initParaState(el);
  el.addEventListener('focus',()=>selectPara(el,false));
  el.addEventListener('blur',()=>savePara(el));
  el.addEventListener('mouseup',saveSelection);
  el.addEventListener('keyup',saveSelection);
});
document.querySelectorAll('.cmtmark').forEach(el=>{
  el.addEventListener('click',e=>{e.stopPropagation();notify({type:'doc-cmt',cid:el.dataset.cid})});
});
bindFmtBar();
notify({type:'doc-ready'});
"""


def _EditorCacheKey(sdocid, stheme):
    scurrent = os.path.join(DocDir(sdocid), "current.docx")
    scomments = os.path.join(DocDir(sdocid), "comments.json")
    if not os.path.isfile(scurrent):
        return ""
    nver = (
        int(os.path.getmtime(scurrent) * 1000),
        os.path.getsize(scurrent),
        int(os.path.getmtime(scomments) * 1000) if os.path.isfile(scomments) else 0,
    )
    return "%s:%s:%s:fmt3" % (sdocid, stheme, nver)


def _TouchEditorCache(skey, shtml):
    global _oEditorHtmlCache
    if skey in _oEditorHtmlCache:
        _oEditorHtmlCache.pop(skey, None)
    _oEditorHtmlCache[skey] = shtml
    while len(_oEditorHtmlCache) > _nEditorCacheMax:
        _oEditorHtmlCache.pop(next(iter(_oEditorHtmlCache)))


def RenderEditorHtml(sdocid, stheme="girly"):
    from docx import Document
    sdir = DocDir(sdocid)
    scurrent = os.path.join(sdir, "current.docx")
    if not os.path.isfile(scurrent):
        raise ValueError("文档不存在")
    vcomments = ReadJson(os.path.join(sdir, "comments.json"), {"items": []}).get("items", [])
    odoc = Document(scurrent)
    oimgurls = _LoadImageUrls(sdocid, scurrent)
    othemes = _LoadThemeFonts(scurrent)
    sbody = _BuildDocumentBodyHtml(
        odoc, oimgurls,
        _PparaCommentsMap(vcomments),
        _PparaCommentsMap(vcomments, "pending"),
        True, othemes,
    )
    sdocid_js = json.dumps(sdocid)
    scss = GetEditorCss(stheme)
    sjs = _EDITOR_JS % {"docid": sdocid_js}
    sfmtbar = (
        '<div class="fmtshell">'
        '<div class="fmtbar">'
        '<div class="grp">'
        '<button type="button" class="fmtbtn" id="fmt_bold" title="加粗">B</button>'
        '<button type="button" class="fmtbtn i" id="fmt_italic" title="倾斜">I</button>'
        '<button type="button" class="fmtbtn u" id="fmt_underline" title="下划线">U</button>'
        '</div>'
        '<div class="grp">'
        '<select class="fmtselect" id="fmt_font" title="字体">'
        '<option value="">字体</option>'
        '<option value="SimSun">宋体</option>'
        '<option value="SimHei">黑体</option>'
        '<option value="KaiTi">楷体</option>'
        '<option value="FangSong">仿宋</option>'
        '<option value="Microsoft YaHei">微软雅黑</option>'
        '<option value="PingFang SC">苹方</option>'
        '<option value="Times New Roman">Times</option>'
        '<option value="Arial">Arial</option>'
        '</select>'
        '<select class="fmtselect" id="fmt_size" title="字号">'
        '<option value="">字号</option>'
        '<option value="10">10pt</option><option value="11">11pt</option>'
        '<option value="12">12pt</option><option value="14">14pt</option>'
        '<option value="16">16pt</option><option value="18">18pt</option>'
        '<option value="20">20pt</option><option value="22">22pt</option>'
        '<option value="24">24pt</option><option value="28">28pt</option>'
        '<option value="32">32pt</option>'
        '</select>'
        '<input type="color" class="fmtcolor" id="fmt_color" value="#4a3f47" title="文字颜色">'
        '</div>'
        '</div></div>'
        '<div class="fmtpeek" title="悬停展开格式工具栏"></div>'
    )
    return (
        '<!DOCTYPE html><html><head><meta charset="utf-8">'
        '<style>' + scss + '</style></head><body>'
        + sfmtbar
        + '<div class="docworkspace"><div class="docpage docpreview">' + sbody + '</div></div>'
        '<div id="docstatus" class="docstatus"></div>'
        '<script>' + sjs + '</script></body></html>'
    )


def InvalidateEditorCache(sdocid=None):
    global _oEditorHtmlCache
    if sdocid is None:
        _oEditorHtmlCache = {}
        return
    vkeys = [k for k in _oEditorHtmlCache if k.startswith(sdocid + ":")]
    for skey in vkeys:
        _oEditorHtmlCache.pop(skey, None)


def GetEditorHtml(sdocid, stheme="girly"):
    spath = os.path.join(DocDir(sdocid), "current.docx")
    if not os.path.isfile(spath):
        raise ValueError("文档不存在")
    if stheme not in _EDITOR_THEME_CSS:
        stheme = "girly"
    scache_key = _EditorCacheKey(sdocid, stheme)
    if scache_key and scache_key in _oEditorHtmlCache:
        return _oEditorHtmlCache[scache_key]
    shtml = RenderEditorHtml(sdocid, stheme)
    if scache_key:
        _TouchEditorCache(scache_key, shtml)
    return shtml


def GetPreviewHtml(sdocid):
    spath = os.path.join(DocDir(sdocid), "preview.html")
    if not os.path.isfile(spath):
        ExtractComments(sdocid)
        spath = os.path.join(DocDir(sdocid), "preview.html")
    with open(spath, "r", encoding="utf-8") as f:
        return f.read()


def MarkTodoDone(sdocid, stodo_id, bdone=True):
    sdir = DocDir(sdocid)
    otodos = ReadJson(os.path.join(sdir, "todos.json"), {"items": []})
    ocomments = ReadJson(os.path.join(sdir, "comments.json"), {"items": []})
    sstatus = "done" if bdone else "pending"
    scid = None
    for oitem in otodos.get("items", []):
        if oitem["id"] == stodo_id:
            oitem["status"] = sstatus
            scid = oitem.get("comment_id")
            break
    for oc in ocomments.get("items", []):
        if scid and oc["id"] == scid:
            oc["status"] = sstatus
    WriteJson(os.path.join(sdir, "todos.json"), otodos)
    WriteJson(os.path.join(sdir, "comments.json"), ocomments)
    return {"progress": _CalcProgress(otodos)}


def ApplyEdit(sdocid, npara_index, snew_text, scomment_id=None, shtml=None, spara_style=None):
    from docx import Document
    scurrent = os.path.join(DocDir(sdocid), "current.docx")
    odoc = Document(scurrent)
    if npara_index < 0 or npara_index >= len(odoc.paragraphs):
        raise ValueError("段落索引无效")
    opara = odoc.paragraphs[npara_index]
    if spara_style:
        _ApplyParaLayout(opara, spara_style)
    if shtml and shtml.strip():
        _ApplyRichHtmlToPara(opara, shtml)
        stext = _SanitizeParaText(opara.text or "")
    else:
        stext = _SanitizeParaText(snew_text)
        if _ParaHasDrawing(opara):
            _SetParaTextPreserveMedia(opara, stext)
        else:
            opara.text = stext
    odoc.save(scurrent)
    InvalidateEditorCache(sdocid)
    BuildPreview(sdocid)
    ometa = ReadJson(os.path.join(DocDir(sdocid), "meta.json"))
    ometa["updated"] = datetime.now().strftime("%Y-%m-%d %H:%M")
    WriteJson(os.path.join(DocDir(sdocid), "meta.json"), ometa)
    if scomment_id:
        MarkTodoDone(sdocid, "todo-%s" % scomment_id, True)
    return {"para_index": npara_index}


_SNAPSHOT_FILES = ("current.docx", "comments.json", "todos.json")


def _StashDir(sdocid):
    return os.path.join(DocDir(sdocid), "revisions", "_stash")


def HasWorkingStash(sdocid):
    return os.path.isfile(os.path.join(_StashDir(sdocid), "current.docx"))


def _CopyDocState(sfrom, sto):
    os.makedirs(sto, exist_ok=True)
    for sname in _SNAPSHOT_FILES:
        spath = os.path.join(sfrom, sname)
        if os.path.isfile(spath):
            shutil.copy2(spath, os.path.join(sto, sname))


def _BodyParaTexts(sdocx_path):
    from docx import Document
    if not os.path.isfile(sdocx_path):
        return []
    odoc = Document(sdocx_path)
    return [opara.text or "" for opara in odoc.paragraphs]


def _BodyParaComparable(sdocx_path):
    from docx import Document
    if not os.path.isfile(sdocx_path):
        return []
    odoc = Document(sdocx_path)
    othemes = _LoadThemeFonts(sdocx_path)
    return [_ParaComparable(opara, othemes) for opara in odoc.paragraphs]


def _DiffParaTexts(vold_plain, vnew_plain, vold_cmp=None, vnew_cmp=None):
    if vold_cmp is None:
        vold_cmp = vold_plain
    if vnew_cmp is None:
        vnew_cmp = vnew_plain
    nmax = max(len(vold_cmp), len(vnew_cmp))
    vchanges = []
    for i in range(nmax):
        sold_cmp = vold_cmp[i] if i < len(vold_cmp) else ""
        snew_cmp = vnew_cmp[i] if i < len(vnew_cmp) else ""
        if sold_cmp != snew_cmp:
            vchanges.append({
                "para_index": i,
                "old": (vold_plain[i] if i < len(vold_plain) else "")[:800],
                "new": (vnew_plain[i] if i < len(vnew_plain) else "")[:800],
            })
    return vchanges


def _DiffTodos(vold_items, vnew_items):
    oold = {x.get("comment_id"): x.get("status") for x in vold_items}
    onew = {x.get("comment_id"): x.get("status") for x in vnew_items}
    vchanges = []
    for scid in sorted(set(oold) | set(onew)):
        if oold.get(scid) != onew.get(scid):
            vchanges.append({
                "comment_id": scid,
                "old": oold.get(scid, ""),
                "new": onew.get(scid, ""),
            })
    return vchanges


def _RevisionParentId(sdocid, srevid):
    vrevs = _LoadRevisionList(sdocid)
    slog = os.path.join(DocDir(sdocid), "revisions", srevid, "log.json")
    if not os.path.isfile(slog):
        return vrevs[0].get("id", "") if vrevs else ""
    for i, orev in enumerate(vrevs):
        if orev.get("id") == srevid and i + 1 < len(vrevs):
            return vrevs[i + 1].get("id", "")
    return ""


def _RevisionParentDoc(sdocid, srevid):
    sdir = DocDir(sdocid)
    sparent_id = _RevisionParentId(sdocid, srevid)
    if sparent_id:
        sparent = os.path.join(sdir, "revisions", sparent_id, "current.docx")
        if os.path.isfile(sparent):
            return sparent
    soriginal = os.path.join(sdir, "original.docx")
    return soriginal if os.path.isfile(soriginal) else ""


def _RevisionParentTodos(sdocid, srevid):
    sparent_id = _RevisionParentId(sdocid, srevid)
    if not sparent_id:
        return []
    spath = os.path.join(DocDir(sdocid), "revisions", sparent_id, "todos.json")
    if os.path.isfile(spath):
        return ReadJson(spath, {"items": []}).get("items", [])
    return []


def _RevHash(srevid):
    return (srevid or "")[-8:]


def _LoadRevisionList(sdocid):
    sdir = DocDir(sdocid)
    ometa = ReadJson(os.path.join(sdir, "meta.json"), {})
    vrevs = list(ometa.get("revisions", []))
    srevroot = os.path.join(sdir, "revisions")
    if os.path.isdir(srevroot):
        for sname in sorted(os.listdir(srevroot), reverse=True):
            if sname.startswith("_"):
                continue
            slog = os.path.join(srevroot, sname, "log.json")
            if os.path.isfile(slog):
                olog = ReadJson(slog)
                if not any(r.get("id") == olog.get("id") for r in vrevs):
                    vrevs.append({
                        "id": olog.get("id"),
                        "hash": olog.get("hash", _RevHash(olog.get("id", ""))),
                        "parent_id": olog.get("parent_id", ""),
                        "parent_hash": olog.get("parent_hash", _RevHash(olog.get("parent_id", "")) if olog.get("parent_id") else ""),
                        "time": olog.get("time"),
                        "message": olog.get("message"),
                        "todos_done": olog.get("todos_done"),
                        "todos_total": olog.get("todos_total"),
                        "progress_percent": olog.get("progress_percent"),
                        "para_change_count": olog.get("para_change_count", len(olog.get("para_changes", []))),
                        "todo_change_count": olog.get("todo_change_count", len(olog.get("todo_changes", []))),
                    })
    vrevs.sort(key=lambda x: x.get("id", ""), reverse=True)
    return vrevs


def _HeadRevisionId(sdocid):
    vrevs = _LoadRevisionList(sdocid)
    return vrevs[0].get("id", "") if vrevs else ""


def _HeadSnapshotPath(sdocid, sname="current.docx"):
    shead = _HeadRevisionId(sdocid)
    sdir = DocDir(sdocid)
    if shead:
        spath = os.path.join(sdir, "revisions", shead, sname)
        if os.path.isfile(spath):
            return spath
    if sname == "current.docx":
        sorig = os.path.join(sdir, "original.docx")
        return sorig if os.path.isfile(sorig) else ""
    return ""


def _DiffStatePaths(sdoc_old, sdoc_new, stodos_old, stodos_new):
    vpara = _DiffParaTexts(
        _BodyParaTexts(sdoc_old), _BodyParaTexts(sdoc_new),
        _BodyParaComparable(sdoc_old), _BodyParaComparable(sdoc_new),
    )
    otodos_old = ReadJson(stodos_old, {"items": []}).get("items", []) if stodos_old and os.path.isfile(stodos_old) else []
    otodos_new = ReadJson(stodos_new, {"items": []}).get("items", []) if stodos_new and os.path.isfile(stodos_new) else []
    vtodos = _DiffTodos(otodos_old, otodos_new)
    return {
        "para_changes": vpara,
        "todo_changes": vtodos,
        "is_dirty": bool(vpara or vtodos),
        "para_change_count": len(vpara),
        "todo_change_count": len(vtodos),
    }


def _FileDigest(spath):
    if not spath or not os.path.isfile(spath):
        return ""
    ohash = hashlib.md5()
    with open(spath, "rb") as f:
        for bchunk in iter(lambda: f.read(65536), b""):
            ohash.update(bchunk)
    return ohash.hexdigest()


def _QuickIsDirty(sdocid):
    sdir = DocDir(sdocid)
    scurrent = os.path.join(sdir, "current.docx")
    shead_doc = _HeadSnapshotPath(sdocid, "current.docx")
    if not os.path.isfile(scurrent) or not shead_doc or not os.path.isfile(shead_doc):
        return False
    if _FileDigest(scurrent) != _FileDigest(shead_doc):
        return True
    stodos_cur = os.path.join(sdir, "todos.json")
    shead_todos = _HeadSnapshotPath(sdocid, "todos.json")
    if shead_todos and os.path.isfile(stodos_cur) and os.path.isfile(shead_todos):
        if _FileDigest(stodos_cur) != _FileDigest(shead_todos):
            return True
    return False


def _QuickWorkingStatus(sdocid):
    bdirty = _QuickIsDirty(sdocid)
    return {
        "para_changes": [],
        "todo_changes": [],
        "is_dirty": bdirty,
        "para_change_count": -1 if bdirty else 0,
        "todo_change_count": 0,
    }


def _DiffWorkingState(sdocid):
    sdir = DocDir(sdocid)
    return _DiffStatePaths(
        _HeadSnapshotPath(sdocid, "current.docx"),
        os.path.join(sdir, "current.docx"),
        _HeadSnapshotPath(sdocid, "todos.json"),
        os.path.join(sdir, "todos.json"),
    )


def GetWorkingStatus(sdocid, bfull=True):
    odiff = _DiffWorkingState(sdocid) if bfull else _QuickWorkingStatus(sdocid)
    vrevs = _LoadRevisionList(sdocid)
    ohead = None
    if vrevs:
        ohead = dict(vrevs[0])
        if not ohead.get("hash"):
            ohead["hash"] = _RevHash(ohead.get("id", ""))
    ostash = {}
    if HasWorkingStash(sdocid):
        ostash = ReadJson(os.path.join(_StashDir(sdocid), "stash.json"), {})
    return {
        "is_dirty": odiff["is_dirty"],
        "para_change_count": odiff["para_change_count"],
        "todo_change_count": odiff["todo_change_count"],
        "para_changes": odiff["para_changes"],
        "todo_changes": odiff["todo_changes"],
        "head": ohead,
        "commit_count": len(vrevs),
        "has_working_stash": HasWorkingStash(sdocid),
        "stash": ostash,
        "baseline": "head" if vrevs else "original",
    }


def DiscardWorkingChanges(sdocid):
    sdir = DocDir(sdocid)
    shead = _HeadRevisionId(sdocid)
    if shead:
        _CopyDocState(os.path.join(sdir, "revisions", shead), sdir)
        BuildPreview(sdocid)
    else:
        sorig = os.path.join(sdir, "original.docx")
        if os.path.isfile(sorig):
            shutil.copy2(sorig, os.path.join(sdir, "current.docx"))
        ExtractComments(sdocid)
    InvalidateEditorCache(sdocid)
    ometa = ReadJson(os.path.join(sdir, "meta.json"))
    ometa["updated"] = datetime.now().strftime("%Y-%m-%d %H:%M")
    WriteJson(os.path.join(sdir, "meta.json"), ometa)
    return GetWorkingStatus(sdocid)


def _ResolveRevPaths(sdocid, srev):
    sdir = DocDir(sdocid)
    if srev in ("HEAD", "head"):
        shead = _HeadRevisionId(sdocid)
        srev = shead if shead else "original"
    if srev in ("WORKING", "working"):
        return (
            os.path.join(sdir, "current.docx"),
            os.path.join(sdir, "todos.json"),
        )
    if srev == "original":
        return (
            os.path.join(sdir, "original.docx"),
            "",
        )
    srevdir = os.path.join(sdir, "revisions", srev)
    return (
        os.path.join(srevdir, "current.docx"),
        os.path.join(srevdir, "todos.json"),
    )


def CompareRevisions(sdocid, srev_a, srev_b):
    sdoc_a, stodos_a = _ResolveRevPaths(sdocid, srev_a)
    sdoc_b, stodos_b = _ResolveRevPaths(sdocid, srev_b)
    if not os.path.isfile(sdoc_a) or not os.path.isfile(sdoc_b):
        raise ValueError("对比版本不存在")
    odiff = _DiffStatePaths(sdoc_b, sdoc_a, stodos_b, stodos_a)
    return {
        "rev_a": srev_a,
        "rev_b": srev_b,
        "hash_a": _RevHash(srev_a) if srev_a not in ("WORKING", "working", "original") else srev_a,
        "hash_b": _RevHash(srev_b) if srev_b not in ("WORKING", "working", "original") else srev_b,
        **odiff,
    }


def _BuildRevisionDiff(sdocid, srevid):
    srevdir = os.path.join(DocDir(sdocid), "revisions", srevid)
    scurrent = os.path.join(srevdir, "current.docx")
    sparent = _RevisionParentDoc(sdocid, srevid)
    vpara = _DiffParaTexts(
        _BodyParaTexts(sparent), _BodyParaTexts(scurrent),
        _BodyParaComparable(sparent), _BodyParaComparable(scurrent),
    )
    otodos_new = ReadJson(os.path.join(srevdir, "todos.json"), {"items": []}).get("items", [])
    otodos_old = _RevisionParentTodos(sdocid, srevid)
    vtodos = _DiffTodos(otodos_old, otodos_new)
    return {"para_changes": vpara, "todo_changes": vtodos}


def SaveRevision(sdocid, smessage):
    smessage = (smessage or "").strip()
    if not smessage:
        raise ValueError("请填写 commit 说明")
    sdir = DocDir(sdocid)
    scurrent = os.path.join(sdir, "current.docx")
    if not os.path.isfile(scurrent):
        raise ValueError("文档不存在")
    srevid = datetime.now().strftime("%Y%m%d-%H%M%S-%f")
    srevdir = os.path.join(sdir, "revisions", srevid)
    os.makedirs(srevdir, exist_ok=True)
    _CopyDocState(sdir, srevdir)
    slog = os.path.join(srevdir, "log.json")
    if os.path.isfile(slog):
        os.remove(slog)
    otodos = ReadJson(os.path.join(sdir, "todos.json"), {"items": []})
    oprog = _CalcProgress(otodos)
    odiff = _BuildRevisionDiff(sdocid, srevid)
    sparent_id = _RevisionParentId(sdocid, srevid)
    orevision = {
        "id": srevid,
        "hash": _RevHash(srevid),
        "parent_id": sparent_id,
        "parent_hash": _RevHash(sparent_id) if sparent_id else "",
        "time": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "message": smessage,
        "todos_done": oprog["done"],
        "todos_total": oprog["total"],
        "progress_percent": oprog["percent"],
        "para_change_count": len(odiff.get("para_changes", [])),
        "todo_change_count": len(odiff.get("todo_changes", [])),
        "para_changes": odiff.get("para_changes", []),
        "todo_changes": odiff.get("todo_changes", []),
    }
    WriteJson(os.path.join(srevdir, "log.json"), orevision)
    ometa = ReadJson(os.path.join(sdir, "meta.json"))
    vrevs = ometa.get("revisions", [])
    vrevs.insert(0, {
        "id": srevid,
        "hash": orevision["hash"],
        "parent_id": sparent_id,
        "parent_hash": orevision["parent_hash"],
        "time": orevision["time"],
        "message": smessage,
        "todos_done": oprog["done"],
        "todos_total": oprog["total"],
        "progress_percent": oprog["percent"],
        "para_change_count": orevision["para_change_count"],
        "todo_change_count": orevision["todo_change_count"],
    })
    ometa["revisions"] = vrevs[:50]
    ometa["updated"] = orevision["time"]
    WriteJson(os.path.join(sdir, "meta.json"), ometa)
    return orevision


def ListRevisions(sdocid):
    vrevs = _LoadRevisionList(sdocid)
    ostatus = GetWorkingStatus(sdocid)
    return {
        "revisions": vrevs,
        "has_working_stash": HasWorkingStash(sdocid),
        "head": ostatus.get("head"),
        "is_dirty": ostatus.get("is_dirty"),
        "working_status": ostatus,
    }


def GetRevisionDetail(sdocid, srevid):
    srevdir = os.path.join(DocDir(sdocid), "revisions", srevid)
    if not os.path.isdir(srevdir) or srevid.startswith("_"):
        raise ValueError("提交记录不存在")
    olog = ReadJson(os.path.join(srevdir, "log.json"))
    if "para_changes" not in olog:
        odiff = _BuildRevisionDiff(sdocid, srevid)
        olog.update(odiff)
    ocomments = ReadJson(os.path.join(srevdir, "comments.json"), {"items": []})
    otodos = ReadJson(os.path.join(srevdir, "todos.json"), {"items": []})
    return {
        "log": olog,
        "comments": ocomments.get("items", []),
        "todos": otodos.get("items", []),
        "has_file": os.path.isfile(os.path.join(srevdir, "current.docx")),
        "has_working_stash": HasWorkingStash(sdocid),
    }


def RestoreRevision(sdocid, srevid):
    sdir = DocDir(sdocid)
    srevdir = os.path.join(sdir, "revisions", srevid)
    if not os.path.isdir(srevdir):
        raise ValueError("提交记录不存在")
    sstash = _StashDir(sdocid)
    if not HasWorkingStash(sdocid):
        _CopyDocState(sdir, sstash)
        WriteJson(os.path.join(sstash, "stash.json"), {
            "time": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "from": "before_restore",
            "target_rev": srevid,
        })
    _CopyDocState(srevdir, sdir)
    InvalidateEditorCache(sdocid)
    BuildPreview(sdocid)
    ometa = ReadJson(os.path.join(sdir, "meta.json"))
    ometa["updated"] = datetime.now().strftime("%Y-%m-%d %H:%M")
    WriteJson(os.path.join(sdir, "meta.json"), ometa)
    return {"id": srevid, "has_working_stash": True}


def RestoreWorkingCopy(sdocid):
    sdir = DocDir(sdocid)
    sstash = _StashDir(sdocid)
    if not HasWorkingStash(sdocid):
        raise ValueError("没有可恢复的工作区修改")
    _CopyDocState(sstash, sdir)
    shutil.rmtree(sstash)
    InvalidateEditorCache(sdocid)
    BuildPreview(sdocid)
    ometa = ReadJson(os.path.join(sdir, "meta.json"))
    ometa["updated"] = datetime.now().strftime("%Y-%m-%d %H:%M")
    WriteJson(os.path.join(sdir, "meta.json"), ometa)
    return {"restored": "working", "has_working_stash": False}


def ExportDoc(sdocid, sdest_dir, sfilename):
    scurrent = os.path.join(DocDir(sdocid), "current.docx")
    if not os.path.isfile(scurrent):
        raise ValueError("文档不存在")
    sdest_dir = os.path.expanduser((sdest_dir or "").strip())
    if not sdest_dir or not os.path.isdir(sdest_dir):
        raise ValueError("请选择有效的导出文件夹")
    sfilename = os.path.basename((sfilename or "export.docx").strip())
    if not sfilename.lower().endswith(".docx"):
        sfilename += ".docx"
    sfilename = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", sfilename).strip("._") or "export.docx"
    sdest = os.path.join(sdest_dir, sfilename)
    from docx import Document
    odoc = Document(scurrent)
    _CleanDocxUiArtifacts(odoc)
    odoc.save(sdest)
    return {"path": sdest, "filename": sfilename, "dir": sdest_dir}


def DeleteDoc(sdocid):
    sdir = DocDir(sdocid)
    if os.path.isdir(sdir):
        shutil.rmtree(sdir)
    omanifest = ReadManifest()
    omanifest["docs"] = [d for d in omanifest.get("docs", []) if d["id"] != sdocid]
    WriteManifest(omanifest)
    return {"id": sdocid}
