#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""文档修订版本：stash、diff、commit、restore。"""
import os
import shutil
import hashlib
from datetime import datetime

import doc_paths as dpaths
from doc_paths import DocDir, ReadJson, WriteJson, CalcTodoProgress
from docx_parser import *  # noqa: F403
import doc_editor_html as dhtml

_SNAPSHOT_FILES = ("current.docx", "comments.json", "todos.json")


def _ExtractComments(sdocid):
    import doc_api
    return doc_api.ExtractComments(sdocid)


def _StashDir(sdocid):
    return os.path.join(DocDir(sdocid), "revisions", "_stash")


def HasWorkingStash(sdocid):
    return os.path.isfile(os.path.join(_StashDir(sdocid), "current.docx"))


def _CopyDocState(sfrom, sto):
    os.makedirs(sto, exist_ok=True)
    for sname in _SNAPSHOT_FILES:
        spath = os.path.join(sfrom, sname)
        if os.path.isfile(spath):
            shutil.copy2(spath, os.path.join(sto, sname))


def _BodyParaTexts(sdocx_path):
    from docx import Document
    if not os.path.isfile(sdocx_path):
        return []
    odoc = Document(sdocx_path)
    return [_SanitizeParaText(opara.text or "") for opara in odoc.paragraphs]


def _BodyParaComparable(sdocx_path):
    from docx import Document
    if not os.path.isfile(sdocx_path):
        return []
    odoc = Document(sdocx_path)
    othemes = _LoadThemeFonts(sdocx_path)
    return [_ParaComparable(opara, othemes) for opara in odoc.paragraphs]


def _DiffParaTexts(vold_plain, vnew_plain, vold_cmp=None, vnew_cmp=None):
    if vold_cmp is None:
        vold_cmp = vold_plain
    if vnew_cmp is None:
        vnew_cmp = vnew_plain
    nmax = max(len(vold_cmp), len(vnew_cmp))
    vchanges = []
    for i in range(nmax):
        sold_cmp = vold_cmp[i] if i < len(vold_cmp) else ""
        snew_cmp = vnew_cmp[i] if i < len(vnew_cmp) else ""
        if sold_cmp != snew_cmp:
            vchanges.append({
                "para_index": i,
                "old": (vold_plain[i] if i < len(vold_plain) else "")[:800],
                "new": (vnew_plain[i] if i < len(vnew_plain) else "")[:800],
            })
    return vchanges


def _DiffTodos(vold_items, vnew_items):
    oold = {x.get("comment_id"): x.get("status") for x in vold_items}
    onew = {x.get("comment_id"): x.get("status") for x in vnew_items}
    vchanges = []
    for scid in sorted(set(oold) | set(onew)):
        if oold.get(scid) != onew.get(scid):
            vchanges.append({
                "comment_id": scid,
                "old": oold.get(scid, ""),
                "new": onew.get(scid, ""),
            })
    return vchanges


def _RevisionParentId(sdocid, srevid):
    vrevs = _LoadRevisionList(sdocid)
    slog = os.path.join(DocDir(sdocid), "revisions", srevid, "log.json")
    if not os.path.isfile(slog):
        return vrevs[0].get("id", "") if vrevs else ""
    for i, orev in enumerate(vrevs):
        if orev.get("id") == srevid and i + 1 < len(vrevs):
            return vrevs[i + 1].get("id", "")
    return ""


def _RevisionParentDoc(sdocid, srevid):
    sdir = DocDir(sdocid)
    sparent_id = _RevisionParentId(sdocid, srevid)
    if sparent_id:
        sparent = os.path.join(sdir, "revisions", sparent_id, "current.docx")
        if os.path.isfile(sparent):
            return sparent
    soriginal = os.path.join(sdir, "original.docx")
    return soriginal if os.path.isfile(soriginal) else ""


def _RevisionParentTodos(sdocid, srevid):
    sparent_id = _RevisionParentId(sdocid, srevid)
    if not sparent_id:
        return []
    spath = os.path.join(DocDir(sdocid), "revisions", sparent_id, "todos.json")
    if os.path.isfile(spath):
        return ReadJson(spath, {"items": []}).get("items", [])
    return []


def _RevHash(srevid):
    return (srevid or "")[-8:]


def _LoadRevisionList(sdocid):
    sdir = DocDir(sdocid)
    ometa = ReadJson(os.path.join(sdir, "meta.json"), {})
    vrevs = list(ometa.get("revisions", []))
    srevroot = os.path.join(sdir, "revisions")
    if os.path.isdir(srevroot):
        for sname in sorted(os.listdir(srevroot), reverse=True):
            if sname.startswith("_"):
                continue
            slog = os.path.join(srevroot, sname, "log.json")
            if os.path.isfile(slog):
                olog = ReadJson(slog)
                if not any(r.get("id") == olog.get("id") for r in vrevs):
                    vrevs.append({
                        "id": olog.get("id"),
                        "hash": olog.get("hash", _RevHash(olog.get("id", ""))),
                        "parent_id": olog.get("parent_id", ""),
                        "parent_hash": olog.get("parent_hash", _RevHash(olog.get("parent_id", "")) if olog.get("parent_id") else ""),
                        "time": olog.get("time"),
                        "message": olog.get("message"),
                        "todos_done": olog.get("todos_done"),
                        "todos_total": olog.get("todos_total"),
                        "progress_percent": olog.get("progress_percent"),
                        "para_change_count": olog.get("para_change_count", len(olog.get("para_changes", []))),
                        "todo_change_count": olog.get("todo_change_count", len(olog.get("todo_changes", []))),
                    })
    vrevs.sort(key=lambda x: x.get("id", ""), reverse=True)
    return vrevs


def _HeadRevisionId(sdocid):
    vrevs = _LoadRevisionList(sdocid)
    return vrevs[0].get("id", "") if vrevs else ""


def _HeadSnapshotPath(sdocid, sname="current.docx"):
    shead = _HeadRevisionId(sdocid)
    sdir = DocDir(sdocid)
    if shead:
        spath = os.path.join(sdir, "revisions", shead, sname)
        if os.path.isfile(spath):
            return spath
    if sname == "current.docx":
        sorig = os.path.join(sdir, "original.docx")
        return sorig if os.path.isfile(sorig) else ""
    return ""


def _DiffStatePaths(sdoc_old, sdoc_new, stodos_old, stodos_new):
    vpara = _DiffParaTexts(
        _BodyParaTexts(sdoc_old), _BodyParaTexts(sdoc_new),
        _BodyParaComparable(sdoc_old), _BodyParaComparable(sdoc_new),
    )
    otodos_old = ReadJson(stodos_old, {"items": []}).get("items", []) if stodos_old and os.path.isfile(stodos_old) else []
    otodos_new = ReadJson(stodos_new, {"items": []}).get("items", []) if stodos_new and os.path.isfile(stodos_new) else []
    vtodos = _DiffTodos(otodos_old, otodos_new)
    return {
        "para_changes": vpara,
        "todo_changes": vtodos,
        "is_dirty": bool(vpara or vtodos),
        "para_change_count": len(vpara),
        "todo_change_count": len(vtodos),
    }


def _FileDigest(spath):
    if not spath or not os.path.isfile(spath):
        return ""
    ohash = hashlib.md5()
    with open(spath, "rb") as f:
        for bchunk in iter(lambda: f.read(65536), b""):
            ohash.update(bchunk)
    return ohash.hexdigest()


def _QuickIsDirty(sdocid):
    sdir = DocDir(sdocid)
    scurrent = os.path.join(sdir, "current.docx")
    shead_doc = _HeadSnapshotPath(sdocid, "current.docx")
    if not os.path.isfile(scurrent) or not shead_doc or not os.path.isfile(shead_doc):
        return False
    if _FileDigest(scurrent) != _FileDigest(shead_doc):
        return True
    stodos_cur = os.path.join(sdir, "todos.json")
    shead_todos = _HeadSnapshotPath(sdocid, "todos.json")
    if shead_todos and os.path.isfile(stodos_cur) and os.path.isfile(shead_todos):
        if _FileDigest(stodos_cur) != _FileDigest(shead_todos):
            return True
    return False


def _QuickWorkingStatus(sdocid):
    bdirty = _QuickIsDirty(sdocid)
    return {
        "para_changes": [],
        "todo_changes": [],
        "is_dirty": bdirty,
        "para_change_count": -1 if bdirty else 0,
        "todo_change_count": 0,
    }


def _DiffWorkingState(sdocid):
    sdir = DocDir(sdocid)
    return _DiffStatePaths(
        _HeadSnapshotPath(sdocid, "current.docx"),
        os.path.join(sdir, "current.docx"),
        _HeadSnapshotPath(sdocid, "todos.json"),
        os.path.join(sdir, "todos.json"),
    )


def GetWorkingStatus(sdocid, bfull=True):
    odiff = _DiffWorkingState(sdocid) if bfull else _QuickWorkingStatus(sdocid)
    vrevs = _LoadRevisionList(sdocid)
    ohead = None
    if vrevs:
        ohead = dict(vrevs[0])
        if not ohead.get("hash"):
            ohead["hash"] = _RevHash(ohead.get("id", ""))
    ostash = {}
    if HasWorkingStash(sdocid):
        ostash = ReadJson(os.path.join(_StashDir(sdocid), "stash.json"), {})
    return {
        "is_dirty": odiff["is_dirty"],
        "para_change_count": odiff["para_change_count"],
        "todo_change_count": odiff["todo_change_count"],
        "para_changes": odiff["para_changes"],
        "todo_changes": odiff["todo_changes"],
        "head": ohead,
        "commit_count": len(vrevs),
        "has_working_stash": HasWorkingStash(sdocid),
        "stash": ostash,
        "baseline": "head" if vrevs else "original",
    }


def DiscardWorkingChanges(sdocid):
    sdir = DocDir(sdocid)
    shead = _HeadRevisionId(sdocid)
    if shead:
        _CopyDocState(os.path.join(sdir, "revisions", shead), sdir)
        dhtml.BuildPreview(sdocid)
    else:
        sorig = os.path.join(sdir, "original.docx")
        if os.path.isfile(sorig):
            shutil.copy2(sorig, os.path.join(sdir, "current.docx"))
        _ExtractComments(sdocid)
    dhtml.InvalidateEditorCache(sdocid)
    ometa = ReadJson(os.path.join(sdir, "meta.json"))
    ometa["updated"] = datetime.now().strftime("%Y-%m-%d %H:%M")
    WriteJson(os.path.join(sdir, "meta.json"), ometa)
    return GetWorkingStatus(sdocid)


def _ResolveRevPaths(sdocid, srev):
    sdir = DocDir(sdocid)
    if srev in ("HEAD", "head"):
        shead = _HeadRevisionId(sdocid)
        srev = shead if shead else "original"
    if srev in ("WORKING", "working"):
        return (
            os.path.join(sdir, "current.docx"),
            os.path.join(sdir, "todos.json"),
        )
    if srev == "original":
        return (
            os.path.join(sdir, "original.docx"),
            "",
        )
    srevdir = os.path.join(sdir, "revisions", srev)
    return (
        os.path.join(srevdir, "current.docx"),
        os.path.join(srevdir, "todos.json"),
    )


def CompareRevisions(sdocid, srev_a, srev_b):
    sdoc_a, stodos_a = _ResolveRevPaths(sdocid, srev_a)
    sdoc_b, stodos_b = _ResolveRevPaths(sdocid, srev_b)
    if not os.path.isfile(sdoc_a) or not os.path.isfile(sdoc_b):
        raise ValueError("对比版本不存在")
    odiff = _DiffStatePaths(sdoc_b, sdoc_a, stodos_b, stodos_a)
    return {
        "rev_a": srev_a,
        "rev_b": srev_b,
        "hash_a": _RevHash(srev_a) if srev_a not in ("WORKING", "working", "original") else srev_a,
        "hash_b": _RevHash(srev_b) if srev_b not in ("WORKING", "working", "original") else srev_b,
        **odiff,
    }


def _BuildRevisionDiff(sdocid, srevid):
    srevdir = os.path.join(DocDir(sdocid), "revisions", srevid)
    scurrent = os.path.join(srevdir, "current.docx")
    sparent = _RevisionParentDoc(sdocid, srevid)
    vpara = _DiffParaTexts(
        _BodyParaTexts(sparent), _BodyParaTexts(scurrent),
        _BodyParaComparable(sparent), _BodyParaComparable(scurrent),
    )
    otodos_new = ReadJson(os.path.join(srevdir, "todos.json"), {"items": []}).get("items", [])
    otodos_old = _RevisionParentTodos(sdocid, srevid)
    vtodos = _DiffTodos(otodos_old, otodos_new)
    return {"para_changes": vpara, "todo_changes": vtodos}


def SaveRevision(sdocid, smessage):
    smessage = (smessage or "").strip()
    if not smessage:
        raise ValueError("请填写 commit 说明")
    sdir = DocDir(sdocid)
    scurrent = os.path.join(sdir, "current.docx")
    if not os.path.isfile(scurrent):
        raise ValueError("文档不存在")
    srevid = datetime.now().strftime("%Y%m%d-%H%M%S-%f")
    srevdir = os.path.join(sdir, "revisions", srevid)
    os.makedirs(srevdir, exist_ok=True)
    _CopyDocState(sdir, srevdir)
    slog = os.path.join(srevdir, "log.json")
    if os.path.isfile(slog):
        os.remove(slog)
    otodos = ReadJson(os.path.join(sdir, "todos.json"), {"items": []})
    oprog = CalcTodoProgress(otodos)
    odiff = _BuildRevisionDiff(sdocid, srevid)
    sparent_id = _RevisionParentId(sdocid, srevid)
    orevision = {
        "id": srevid,
        "hash": _RevHash(srevid),
        "parent_id": sparent_id,
        "parent_hash": _RevHash(sparent_id) if sparent_id else "",
        "time": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "message": smessage,
        "todos_done": oprog["done"],
        "todos_total": oprog["total"],
        "progress_percent": oprog["percent"],
        "para_change_count": len(odiff.get("para_changes", [])),
        "todo_change_count": len(odiff.get("todo_changes", [])),
        "para_changes": odiff.get("para_changes", []),
        "todo_changes": odiff.get("todo_changes", []),
    }
    WriteJson(os.path.join(srevdir, "log.json"), orevision)
    ometa = ReadJson(os.path.join(sdir, "meta.json"))
    vrevs = ometa.get("revisions", [])
    vrevs.insert(0, {
        "id": srevid,
        "hash": orevision["hash"],
        "parent_id": sparent_id,
        "parent_hash": orevision["parent_hash"],
        "time": orevision["time"],
        "message": smessage,
        "todos_done": oprog["done"],
        "todos_total": oprog["total"],
        "progress_percent": oprog["percent"],
        "para_change_count": orevision["para_change_count"],
        "todo_change_count": orevision["todo_change_count"],
    })
    ometa["revisions"] = vrevs[:50]
    ometa["updated"] = orevision["time"]
    WriteJson(os.path.join(sdir, "meta.json"), ometa)
    return orevision


def ListRevisions(sdocid):
    vrevs = _LoadRevisionList(sdocid)
    ostatus = GetWorkingStatus(sdocid)
    return {
        "revisions": vrevs,
        "has_working_stash": HasWorkingStash(sdocid),
        "head": ostatus.get("head"),
        "is_dirty": ostatus.get("is_dirty"),
        "working_status": ostatus,
    }


def GetRevisionDetail(sdocid, srevid):
    srevdir = os.path.join(DocDir(sdocid), "revisions", srevid)
    if not os.path.isdir(srevdir) or srevid.startswith("_"):
        raise ValueError("提交记录不存在")
    olog = ReadJson(os.path.join(srevdir, "log.json"))
    if "para_changes" not in olog:
        odiff = _BuildRevisionDiff(sdocid, srevid)
        olog.update(odiff)
    ocomments = ReadJson(os.path.join(srevdir, "comments.json"), {"items": []})
    otodos = ReadJson(os.path.join(srevdir, "todos.json"), {"items": []})
    return {
        "log": olog,
        "comments": ocomments.get("items", []),
        "todos": otodos.get("items", []),
        "has_file": os.path.isfile(os.path.join(srevdir, "current.docx")),
        "has_working_stash": HasWorkingStash(sdocid),
    }


def RestoreRevision(sdocid, srevid):
    sdir = DocDir(sdocid)
    srevdir = os.path.join(sdir, "revisions", srevid)
    if not os.path.isdir(srevdir):
        raise ValueError("提交记录不存在")
    sstash = _StashDir(sdocid)
    if not HasWorkingStash(sdocid):
        _CopyDocState(sdir, sstash)
        WriteJson(os.path.join(sstash, "stash.json"), {
            "time": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "from": "before_restore",
            "target_rev": srevid,
        })
    _CopyDocState(srevdir, sdir)
    dhtml.InvalidateEditorCache(sdocid)
    dhtml.BuildPreview(sdocid)
    ometa = ReadJson(os.path.join(sdir, "meta.json"))
    ometa["updated"] = datetime.now().strftime("%Y-%m-%d %H:%M")
    WriteJson(os.path.join(sdir, "meta.json"), ometa)
    return {"id": srevid, "has_working_stash": True}


def RestoreWorkingCopy(sdocid):
    sdir = DocDir(sdocid)
    sstash = _StashDir(sdocid)
    if not HasWorkingStash(sdocid):
        raise ValueError("没有可恢复的工作区修改")
    _CopyDocState(sstash, sdir)
    shutil.rmtree(sstash)
    dhtml.InvalidateEditorCache(sdocid)
    dhtml.BuildPreview(sdocid)
    ometa = ReadJson(os.path.join(sdir, "meta.json"))
    ometa["updated"] = datetime.now().strftime("%Y-%m-%d %H:%M")
    WriteJson(os.path.join(sdir, "meta.json"), ometa)
    return {"restored": "working", "has_working_stash": False}
